import csv
import re
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree

from bs4 import BeautifulSoup, Comment
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from markdownify import markdownify as md
from pypdf import PdfReader

from app.converters.base import ConversionResult
from app.converters.web_extractors.utils import clean_markdown, normalize_text


MAX_TEXT_TITLE_LENGTH = 90
HTML_CONTENT_SELECTORS = [
    "article",
    "main article",
    "main",
    "[role='main']",
    ".markdown-body",
    ".prose",
    ".entry-content",
    ".post-content",
    ".article-content",
    ".content",
    "#content",
    "#main",
]
HTML_NOISE_RE = re.compile(
    r"(advert|banner|breadcrumb|cookie|footer|header|login|menu|modal|nav|newsletter|"
    r"pagination|popup|promo|search|share|sidebar|signin|signup|subscribe|toolbar)",
    re.I,
)


@dataclass
class DocumentStats:
    paragraphs: int = 0
    headings: int = 0
    tables: int = 0
    rows: int = 0
    images: int = 0
    formulas: int = 0
    pages: int = 0
    characters: int = 0


@dataclass
class PdfDocumentInfo:
    title: str
    author: str | None = None
    source_url: str | None = None
    created_at: str | None = None


@dataclass
class PdfTextBlock:
    page: int
    text: str
    bbox: tuple[float, float, float, float]
    font_size: float
    is_bold: bool = False
    is_formula: bool = False
    is_image: bool = False
    image_path: str | None = None
    is_table: bool = False


@dataclass
class PdfExtraction:
    body: str
    resources: list[str]
    stats: DocumentStats
    metadata: dict


@dataclass
class PdfLineSegment:
    orientation: str
    position: float
    start: float
    end: float


def convert_text(path: Path, source_type: str = "text") -> ConversionResult:
    text = normalize_plain_text(path.read_text(encoding="utf-8", errors="ignore"))
    title = infer_plain_text_title(text, path.stem or "文本内容")
    body = clean_markdown(text, title)
    return ConversionResult(
        title=title,
        source_type=source_type,
        body=body,
        summary_seed=body,
        metadata=stats_metadata("text", DocumentStats(paragraphs=count_paragraphs(body), characters=len(body))),
    )


def convert_html(path: Path) -> ConversionResult:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    title = html_title(soup, path.stem)
    clean_html_soup(soup)
    node = select_html_content_node(soup)
    strip_internal_html_attrs(node)
    body = md(
        str(node),
        heading_style="ATX",
        bullets="-",
        strip=["script", "style"],
    )
    body = clean_markdown(body, title)
    stats = DocumentStats(
        paragraphs=count_paragraphs(body),
        headings=count_markdown_headings(body),
        tables=count_markdown_tables(body),
        images=len(re.findall(r"!\[[^\]]*]\([^)]+\)", body)),
        characters=len(body),
    )
    return ConversionResult(title=title, source_type="html", body=body, summary_seed=body, metadata=stats_metadata("html", stats))


def convert_csv(path: Path) -> ConversionResult:
    rows = read_csv_rows(path)
    if not rows:
        body = ""
        stats = DocumentStats()
    else:
        body = markdown_table(rows)
        stats = DocumentStats(tables=1, rows=max(len(rows) - 1, 0), characters=len(body))
    return ConversionResult(title=path.stem, source_type="csv", body=body, summary_seed=body, metadata=stats_metadata("csv", stats))


def convert_pdf(path: Path, assets_dir: Path | None = None) -> ConversionResult:
    reader = PdfReader(str(path))
    info = pdf_document_info(reader, path)
    if assets_dir is not None:
        assets_dir.mkdir(parents=True, exist_ok=True)

    try:
        extraction = extract_pdf_with_pymupdf(path, info, assets_dir)
    except Exception as exc:
        extraction = extract_pdf_with_pypdf(reader, info)
        extraction.metadata["pdf_engine_fallback_reason"] = str(exc)

    body = clean_markdown(extraction.body, info.title)
    metadata = stats_metadata("pdf", extraction.stats)
    metadata.update(
        {
            **extraction.metadata,
            "pdf_title": info.title,
            "pdf_author": info.author,
            "pdf_source_url": info.source_url,
        }
    )
    if not body:
        metadata["warning"] = "No extractable text found. Scanned PDFs need OCR, which is planned for a later phase."
    return ConversionResult(
        title=info.title,
        source_type="pdf",
        body=body,
        summary_seed=body,
        resources=extraction.resources,
        author=info.author,
        source_url=info.source_url,
        created_at=info.created_at,
        metadata=metadata,
    )


def convert_docx(path: Path, assets_dir: Path) -> ConversionResult:
    assets_dir.mkdir(parents=True, exist_ok=True)
    document = Document(str(path))
    parts: list[str] = []
    stats = DocumentStats()
    title = path.stem
    image_map = extract_docx_images(document, assets_dir)
    emitted_images: set[str] = set()

    for block in iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            markdown = paragraph_markdown(block)
            image_ids = paragraph_image_relationship_ids(block)
            if not parts and markdown.startswith("# "):
                title = markdown[2:].strip() or title
            if markdown:
                parts.append(markdown)
                stats.paragraphs += 1
                if markdown.startswith("#"):
                    stats.headings += 1
            for rel_id in image_ids:
                image_path = image_map.get(rel_id)
                if image_path and image_path not in emitted_images:
                    alt = Path(image_path).stem
                    parts.append(f"![{alt}]({image_path})")
                    emitted_images.add(image_path)
        elif isinstance(block, Table):
            rows = table_rows(block)
            if rows:
                parts.append(markdown_table(rows))
                stats.tables += 1
                stats.rows += max(len(rows) - 1, 0)

    resources = [f"- 图片资源：{path}" for path in image_map.values()]
    stats.images = len(image_map)
    body = clean_markdown("\n\n".join(parts), title)
    stats.characters = len(body)
    return ConversionResult(
        title=title,
        source_type="docx",
        body=body,
        summary_seed=body,
        resources=resources,
        metadata=stats_metadata("docx", stats),
    )


def convert_by_extension(path: Path, assets_dir: Path) -> ConversionResult:
    extension = path.suffix.lower()
    if extension in {".txt", ".md"}:
        return convert_text(path)
    if extension in {".html", ".htm"}:
        return convert_html(path)
    if extension == ".csv":
        return convert_csv(path)
    if extension == ".pdf":
        return convert_pdf(path, assets_dir)
    if extension == ".docx":
        return convert_docx(path, assets_dir)
    raise ValueError(f"Unsupported document type: {extension}")


def normalize_plain_text(text: str) -> str:
    text = text.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.splitlines()]
    text = "\n".join(lines)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def infer_plain_text_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        candidate = normalize_text(line.strip("# "))
        if not candidate:
            continue
        if len(candidate) <= MAX_TEXT_TITLE_LENGTH:
            return candidate
        break
    return fallback


def clean_html_soup(soup: BeautifulSoup) -> None:
    for comment in soup.find_all(string=lambda value: isinstance(value, Comment)):
        comment.extract()
    for tag in soup(["script", "style", "noscript", "template", "head", "header", "nav", "footer", "aside", "form", "button", "input"]):
        tag.decompose()
    for tag in list(soup.find_all(True)):
        if not getattr(tag, "name", None):
            continue
        marker = " ".join(
            str(value)
            for value in [
                tag.get("id", ""),
                " ".join(tag.get("class", []) if isinstance(tag.get("class"), list) else [str(tag.get("class", ""))]),
                tag.get("role", ""),
                tag.get("aria-label", ""),
            ]
        )
        if tag.name not in {"article", "main"} and HTML_NOISE_RE.search(marker):
            tag.decompose()
    for tag in soup.find_all(True):
        tag.attrs = {
            key: value
            for key, value in tag.attrs.items()
            if key in {"id", "class", "role", "href", "src", "alt", "title", "datetime", "colspan", "rowspan"}
        }


def select_html_content_node(soup: BeautifulSoup):
    candidates = []
    for selector in HTML_CONTENT_SELECTORS:
        candidates.extend(soup.select(selector))
    if soup.body:
        candidates.append(soup.body)
    else:
        candidates.append(soup)

    best = None
    best_score = float("-inf")
    seen: set[int] = set()
    for node in candidates:
        identity = id(node)
        if identity in seen:
            continue
        seen.add(identity)
        score = html_content_score(node)
        if score > best_score:
            best = node
            best_score = score
    return best or soup.body or soup


def strip_internal_html_attrs(node) -> None:
    for tag in node.find_all(True):
        tag.attrs = {
            key: value
            for key, value in tag.attrs.items()
            if key in {"href", "src", "alt", "title", "datetime", "colspan", "rowspan"}
        }


def html_content_score(node) -> float:
    text = normalize_text(node.get_text(" ", strip=True))
    if not text:
        return -1000
    text_len = len(text)
    paragraph_count = len([p for p in node.find_all("p") if len(normalize_text(p.get_text(" ", strip=True))) >= 20])
    heading_count = len(node.find_all(re.compile("^h[1-6]$")))
    table_count = len(node.find_all("table"))
    code_count = len(node.find_all(["pre", "code"]))
    image_count = len(node.find_all("img"))
    anchor_text = sum(len(normalize_text(a.get_text(" ", strip=True))) for a in node.find_all("a"))
    link_density = anchor_text / max(text_len, 1)
    marker = " ".join(
        str(value)
        for value in [
            node.get("id", ""),
            " ".join(node.get("class", []) if isinstance(node.get("class"), list) else [str(node.get("class", ""))]),
            node.get("role", ""),
        ]
    )
    noise_penalty = 80 if HTML_NOISE_RE.search(marker) else 0
    return (
        min(text_len, 30000) / 80
        + paragraph_count * 16
        + heading_count * 8
        + table_count * 18
        + code_count * 12
        + image_count * 3
        - max(0.0, link_density - 0.35) * 180
        - noise_penalty
    )


def html_title(soup: BeautifulSoup, fallback: str) -> str:
    for selector in ["h1", "title", "[property='og:title']", "[name='twitter:title']"]:
        node = soup.select_one(selector)
        if not node:
            continue
        value = node.get("content") if node.name == "meta" else node.get_text(" ", strip=True)
        if isinstance(value, str) and normalize_text(value):
            return normalize_text(value)
    return fallback


def read_csv_rows(path: Path) -> list[list[str]]:
    sample = path.read_text(encoding="utf-8-sig", errors="ignore")[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except Exception:
        dialect = csv.excel
    with path.open("r", encoding="utf-8-sig", newline="", errors="ignore") as handle:
        return [[normalize_text(cell) for cell in row] for row in csv.reader(handle, dialect)]


def normalize_table_row(row: list[str], width: int) -> list[str]:
    return [normalize_table_cell(cell) for cell in (row + [""] * (width - len(row)))[:width]]


def normalize_table_cell(value: str) -> str:
    value = normalize_text(value.replace("\n", " "))
    value = value.replace("\\", "\\\\").replace("|", "\\|")
    return value


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [normalize_table_row(row, width) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def extract_pdf_with_pymupdf(path: Path, info: PdfDocumentInfo, assets_dir: Path | None = None) -> PdfExtraction:
    import fitz

    document = fitz.open(str(path))
    page_count = document.page_count
    page_blocks: list[list[PdfTextBlock]] = []
    all_blocks: list[PdfTextBlock] = []
    used_image_names: set[str] = set()

    for page_index, page in enumerate(document, start=1):
        blocks = pymupdf_page_blocks(page, page_index, assets_dir, used_image_names)
        page_blocks.append(blocks)
        all_blocks.extend(blocks)

    recurring = recurring_pdf_noise_keys(page_blocks)
    font_sizes = [block.font_size for block in all_blocks if block.text and not block.is_image]
    body_size = median(font_sizes) if font_sizes else 10.0
    inferred_title = infer_pdf_title_from_blocks(page_blocks[0] if page_blocks else [], info.title, body_size)
    if should_replace_pdf_title(info.title, path, inferred_title):
        info.title = inferred_title
    image_resources = [block.image_path for block in all_blocks if block.is_image and block.image_path]

    parts: list[str] = []
    stats = DocumentStats(pages=page_count, images=len(image_resources))
    for page_index, page in enumerate(document, start=1):
        blocks = [
            block
            for block in page_blocks[page_index - 1]
            if not should_drop_pdf_block(block, page.rect.height, recurring)
        ]
        if page_index == 1:
            blocks = strip_pdf_title_blocks(blocks, info.title)
        text = render_pdf_blocks(blocks, body_size)
        if not text:
            continue
        page_table_count = sum(1 for block in blocks if block.is_table)
        stats.characters += len(text)
        stats.paragraphs += count_paragraphs(text)
        stats.headings += count_markdown_headings(text)
        stats.tables += max(count_markdown_tables(text), page_table_count)
        stats.rows += count_markdown_table_rows(text)
        stats.formulas += count_markdown_formulas(text)
        parts.append(f"## Page {page_index}\n\n{text}")

    body = repair_pdf_text_artifacts(repair_pdf_page_continuations("\n\n".join(parts)))
    resources = [f"- 图片资源：{path}" for path in image_resources]
    metadata = {
        "pdf_engine": "pymupdf",
        "pdf_body_font_size": round(body_size, 2),
        "pdf_image_count": len(image_resources),
        "pdf_formula_engine": "symbolic_preserve",
        "pdf_formula_ocr_provider": None,
        "pdf_formula_quality": "layout_text_preserved",
        "pdf_table_engine": "pymupdf_find_tables",
    }
    if not body and page_count > 0:
        metadata["warning"] = "No extractable text found. Scanned PDFs need OCR, which is planned for a later phase."
    document.close()
    return PdfExtraction(body=body, resources=resources, stats=stats, metadata=metadata)


def pymupdf_page_blocks(page, page_number: int, assets_dir: Path | None = None, used_image_names: set[str] | None = None) -> list[PdfTextBlock]:
    raw = page.get_text("dict")
    table_blocks = pdf_table_blocks(page, page_number)
    table_bboxes = [block.bbox for block in table_blocks]
    blocks: list[PdfTextBlock] = list(table_blocks)
    image_index = 0
    for block in raw.get("blocks", []):
        if block.get("type") == 1:
            if bbox_overlaps_any(block.get("bbox", (0, 0, 0, 0)), table_bboxes):
                continue
            image_index += 1
            image_block = pdf_image_block(block, page_number, image_index, assets_dir, used_image_names)
            if image_block:
                blocks.append(image_block)
            continue
        if block.get("type") != 0:
            continue
        if bbox_overlaps_any(block.get("bbox", (0, 0, 0, 0)), table_bboxes):
            continue
        lines: list[str] = []
        sizes: list[float] = []
        fonts: list[str] = []
        math_span_count = 0
        span_count = 0
        for line in block.get("lines", []):
            spans = sorted(line.get("spans", []), key=lambda span: span.get("bbox", (0, 0, 0, 0))[0])
            line_text = join_pdf_spans(spans)
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                span_count += 1
                font = str(span.get("font") or "")
                if is_pdf_math_font(font):
                    math_span_count += 1
                if span.get("size"):
                    sizes.append(float(span["size"]))
                if font:
                    fonts.append(font)
            line_text = normalize_text(line_text)
            if line_text:
                lines.append(line_text)
        text = "\n".join(lines).strip()
        if not text:
            continue
        bbox = tuple(float(value) for value in block.get("bbox", (0, 0, 0, 0)))
        font_size = median(sizes) if sizes else 10.0
        is_bold = any("bold" in font.lower() or "black" in font.lower() for font in fonts)
        math_ratio = math_span_count / max(span_count, 1)
        blocks.append(
            PdfTextBlock(
                page=page_number,
                text=text,
                bbox=bbox,
                font_size=font_size,
                is_bold=is_bold,
                is_formula=is_pdf_formula_block(text, math_ratio),
            )
        )
    return sorted(blocks, key=lambda item: (round(item.bbox[1], 1), round(item.bbox[0], 1)))


def pdf_table_blocks(page, page_number: int) -> list[PdfTextBlock]:
    try:
        finder = page.find_tables()
    except Exception:
        return []
    blocks: list[PdfTextBlock] = []
    pending_header: tuple[list[list[str]], tuple[float, float, float, float]] | None = None
    for table_index, table in enumerate(getattr(finder, "tables", []), start=1):
        try:
            rows = table.extract()
        except Exception:
            continue
        normalized = normalize_pdf_table_rows(rows)
        if is_false_positive_pdf_paragraph_table(normalized):
            pending_header = None
            continue
        bbox = tuple(float(value) for value in getattr(table, "bbox", (0, 0, 0, 0)))
        grid_rows: list[list[str]] = []
        coordinate_rows: list[list[str]] = []
        if pending_header and should_merge_adjacent_pdf_table_regions(pending_header[1], bbox):
            combined_bbox = union_pdf_bboxes(pending_header[1], bbox)
            combined_rows = reconstruct_pdf_table_from_words(page, combined_bbox)
            if is_useful_reconstructed_pdf_table(combined_rows):
                normalized = pending_header[0] + normalized
                bbox = combined_bbox
                coordinate_rows = combined_rows
            pending_header = None
        if not is_useful_pdf_table(normalized) and not coordinate_rows:
            if is_potential_pdf_table_header(normalized):
                pending_header = (normalized, bbox)
            continue
        if not coordinate_rows:
            grid_rows = reconstruct_pdf_grid_table_from_page(page, bbox)
            coordinate_rows = grid_rows or reconstruct_pdf_table_from_words(page, bbox)
        if should_render_pdf_table_as_form(normalized):
            table_text = render_pdf_form_table(normalized)
        elif grid_rows and should_use_grid_pdf_table(normalized, grid_rows):
            table_text = render_coordinate_pdf_table(grid_rows)
        elif should_prefer_extracted_pdf_table(normalized, coordinate_rows):
            table_text = render_pdf_markdown_table(normalized)
        elif should_render_pdf_table_as_layout_block(normalized):
            table_text = render_pdf_table_layout_block(normalized, page=page, bbox=bbox)
        elif should_use_coordinate_pdf_table(normalized, coordinate_rows):
            coordinate_rows = apply_extracted_row_labels_to_coordinate_rows(normalized, coordinate_rows)
            table_text = render_coordinate_pdf_table(coordinate_rows)
        else:
            table_text = render_pdf_markdown_table(normalized)
        blocks.append(
            PdfTextBlock(
                page=page_number,
                text=table_text,
                bbox=bbox,
                font_size=0,
                is_table=True,
            )
        )
    return blocks


def normalize_pdf_table_rows(rows: list[list[str | None]]) -> list[list[str]]:
    normalized: list[list[str]] = []
    for row in rows:
        cells = [normalize_pdf_table_cell(cell) for cell in row]
        if any(cells):
            normalized.append(cells)
    return normalized


def normalize_pdf_table_cell(cell: str | None) -> str:
    raw = str(cell or "").replace("\r\n", "\n").replace("\r", "\n")
    raw = re.sub(r"(?:[ \t]*\.[ \t]*){3,}", " … ", raw)
    lines = []
    for line in raw.split("\n"):
        text = normalize_text(line)
        text = re.sub(r"\s+([:;,.])", r"\1", text)
        if text:
            lines.append(text)
    return "\n".join(lines)


def is_useful_pdf_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False
    width = max((len(row) for row in rows), default=0)
    if width < 2:
        return False
    filled_cells = sum(1 for row in rows for cell in row if cell)
    return filled_cells >= 4


def is_useful_reconstructed_pdf_table(rows: list[list[str]]) -> bool:
    if len(rows) < 3:
        return False
    width = max((len(row) for row in rows), default=0)
    if width < 2 or width > 32:
        return False
    filled_cells = sum(1 for row in rows for cell in row if cell)
    return filled_cells >= 8


def is_potential_pdf_table_header(rows: list[list[str]]) -> bool:
    if len(rows) != 1:
        return False
    row = rows[0]
    if not looks_like_pdf_table_header(row):
        return False
    text = " ".join(flatten_pdf_table_cell(cell) for cell in row if cell)
    return len(text.split()) >= 2


def should_merge_adjacent_pdf_table_regions(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> bool:
    vertical_gap = second[1] - first[3]
    if vertical_gap < -1 or vertical_gap > 10:
        return False
    first_width = max(first[2] - first[0], 1)
    second_width = max(second[2] - second[0], 1)
    width_ratio = min(first_width, second_width) / max(first_width, second_width)
    x_overlap = max(0.0, min(first[2], second[2]) - max(first[0], second[0]))
    overlap_ratio = x_overlap / max(min(first_width, second_width), 1)
    return width_ratio >= 0.80 and overlap_ratio >= 0.80


def union_pdf_bboxes(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> tuple[float, float, float, float]:
    return (
        min(first[0], second[0]),
        min(first[1], second[1]),
        max(first[2], second[2]),
        max(first[3], second[3]),
    )


def compact_pdf_table(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return rows
    width = max(len(row) for row in rows)
    padded = [(row + [""] * (width - len(row)))[:width] for row in rows]
    keep_columns = [index for index in range(width) if any(row[index] for row in padded)]
    compacted = [[row[index] for index in keep_columns] for row in padded]
    result: list[list[str]] = []
    seen: set[tuple[str, ...]] = set()
    for row in compacted:
        key = tuple(row)
        if key in seen:
            continue
        seen.add(key)
        result.append(row)
    return result


def pdf_table_cell_lines(cell: str) -> list[str]:
    lines = [normalize_text(line) for line in str(cell or "").splitlines()]
    return [line for line in lines if line]


def flatten_pdf_table_cell(cell: str, separator: str = " ") -> str:
    return separator.join(pdf_table_cell_lines(cell))


def should_render_pdf_table_as_form(rows: list[list[str]]) -> bool:
    width = max((len(row) for row in rows), default=0)
    if width < 3:
        return False
    total = max(len(rows) * width, 1)
    filled = sum(1 for row in rows for cell in row if cell)
    empty_ratio = 1 - filled / total
    non_empty_per_row = filled / max(len(rows), 1)
    long_cells = sum(1 for row in rows for cell in row if len(cell) >= 80)
    if width <= 4 and long_cells >= 2 and non_empty_per_row <= 3.5:
        return True
    return empty_ratio >= 0.35 and (non_empty_per_row <= 3.5 or long_cells >= 2)


def should_render_pdf_table_as_layout_block(rows: list[list[str]]) -> bool:
    raw_width = max((len(row) for row in rows), default=0)
    if raw_width >= 12 and len(rows) >= 3:
        raw_total = max(raw_width * len(rows), 1)
        raw_filled = sum(1 for row in rows for cell in row if cell)
        raw_empty_ratio = 1 - raw_filled / raw_total
        if raw_empty_ratio >= 0.42:
            return True
    compacted = compact_pdf_table(rows)
    width = max((len(row) for row in compacted), default=0)
    if width == 0 or len(compacted) < 3:
        return False
    total_cells = max(width * len(compacted), 1)
    filled = [cell for row in compacted for cell in row if cell]
    if not filled:
        return False
    empty_ratio = 1 - len(filled) / total_cells
    if width >= 12 and empty_ratio >= 0.42:
        return True
    if table_has_expandable_multiline_rows(compacted):
        return False
    if width > 4:
        return False
    average_tokens = sum(len(cell.split()) for cell in filled) / len(filled)
    numeric_tokens = sum(1 for cell in filled for token in cell.split() if re.search(r"\d", token))
    max_tokens = max(len(cell.split()) for cell in filled)
    scientific_headers = {
        "bleu",
        "drop",
        "model",
        "params",
        "perplexity",
        "ppl",
        "steps",
        "train",
    }
    header_text = " ".join(flatten_pdf_table_cell(cell).lower() for cell in compacted[0])
    has_scientific_header = sum(1 for marker in scientific_headers if marker in header_text) >= 2
    return (average_tokens >= 6 and numeric_tokens >= 6) or max_tokens >= 12 or has_scientific_header


def is_false_positive_pdf_paragraph_table(rows: list[list[str]]) -> bool:
    compacted = compact_pdf_table(rows)
    if len(compacted) < 3:
        return False
    width = max((len(row) for row in compacted), default=0)
    if width > 3:
        return False
    filled_per_row = [sum(1 for cell in row if cell) for row in compacted]
    mostly_single_column = sum(1 for count in filled_per_row if count <= 1) >= max(2, len(filled_per_row) - 1)
    if not mostly_single_column:
        return False
    first_row = compacted[0]
    if looks_like_pdf_table_header(first_row):
        return False
    filled = [flatten_pdf_table_cell(cell) for row in compacted for cell in row if cell]
    if not filled:
        return False
    average_words = sum(len(cell.split()) for cell in filled) / len(filled)
    numeric_ratio = sum(1 for cell in filled if re.search(r"\d", cell)) / len(filled)
    sentence_like = sum(1 for cell in filled if looks_like_pdf_sentence(cell) or cell.endswith((",", ";", ":", "and", "or")))
    return average_words >= 5 and numeric_ratio <= 0.25 and sentence_like >= max(2, len(filled) // 2)


def looks_like_pdf_table_header(row: list[str]) -> bool:
    filled = [flatten_pdf_table_cell(cell) for cell in row if cell]
    if len(filled) < 2:
        return False
    if all(1 <= len(cell.split()) <= 4 and len(cell) <= 40 for cell in filled):
        return True
    header_markers = {
        "accuracy",
        "answer",
        "benchmark",
        "bleu",
        "category",
        "contamination",
        "degradation",
        "english",
        "gpt",
        "language",
        "model",
        "score",
        "swahili",
        "value",
    }
    marker_hits = sum(1 for cell in filled for marker in header_markers if marker in cell.lower())
    return marker_hits >= 2


def should_prefer_extracted_pdf_table(extracted_rows: list[list[str]], coordinate_rows: list[list[str]]) -> bool:
    compacted = compact_pdf_table(extracted_rows)
    if len(compacted) < 2 or not coordinate_rows:
        return False
    extracted_width = max((len(row) for row in compacted), default=0)
    coordinate_width = max((len(row) for row in coordinate_rows), default=0)
    if extracted_width == 0 or coordinate_width < max(extracted_width * 4, 10):
        return False
    if not looks_like_pdf_table_header(compacted[0]):
        return False
    if not table_has_expandable_multiline_rows(compacted):
        return False
    return pdf_table_numeric_cell_ratio(coordinate_rows) < 0.16


def pdf_table_numeric_cell_ratio(rows: list[list[str]]) -> float:
    filled = [cell for row in rows for cell in row if cell]
    if not filled:
        return 0.0
    return sum(1 for cell in filled if re.search(r"\d", cell)) / len(filled)


def render_pdf_table_layout_block(rows: list[list[str]], page=None, bbox: tuple[float, float, float, float] | None = None) -> str:
    compacted = compact_pdf_table(rows)
    lines = ["**Table region (layout preserved)**", "", "```text"]
    layout_lines = pdf_table_layout_lines_from_page(page, bbox) if page is not None and bbox is not None else []
    if layout_lines:
        lines.extend(layout_lines)
    else:
        for row in compacted:
            values = [flatten_pdf_table_cell(cell, " / ") for cell in row if cell]
            if values:
                lines.append(" | ".join(values))
    lines.append("```")
    return "\n".join(lines)


def render_pdf_form_table(rows: list[list[str]]) -> str:
    lines = ["**Form/table region**"]
    seen: set[str] = set()
    for row in rows:
        values = [cell for cell in row if cell]
        if not values:
            continue
        text = " — ".join(flatten_pdf_table_cell(value) for value in values)
        if text in seen:
            continue
        seen.add(text)
        lines.append(f"- {text}")
    return "\n".join(lines)


def render_pdf_markdown_table(rows: list[list[str]]) -> str:
    return markdown_table(expand_multiline_pdf_table_rows(compact_pdf_table(rows)))


def render_coordinate_pdf_table(rows: list[list[str]]) -> str:
    rows = collapse_pdf_spanner_columns(remove_duplicate_pdf_table_rows(rows))
    return markdown_table(compact_pdf_table(rows))


def collapse_pdf_spanner_columns(rows: list[list[str]]) -> list[list[str]]:
    if len(rows) < 2:
        return rows
    width = max(len(row) for row in rows)
    padded = [(row + [""] * (width - len(row)))[:width] for row in rows]
    header = padded[0]
    groups: list[list[int]] = []
    index = 0
    while index < width:
        group = [index]
        if header[index]:
            cursor = index + 1
            while cursor < width and not header[cursor]:
                group.append(cursor)
                cursor += 1
            index = cursor
        else:
            index += 1
        groups.append(group)
    if all(len(group) == 1 for group in groups):
        return rows
    collapsed: list[list[str]] = []
    for row_index, row in enumerate(padded):
        collapsed_row: list[str] = []
        for group in groups:
            if row_index == 0:
                collapsed_row.append(row[group[0]])
                continue
            values = [row[column] for column in group if row[column]]
            collapsed_row.append(" / ".join(values))
        collapsed.append(collapsed_row)
    return collapsed


def should_use_coordinate_pdf_table(extracted_rows: list[list[str]], coordinate_rows: list[list[str]]) -> bool:
    if len(coordinate_rows) < 3:
        return False
    coordinate_width = max((len(row) for row in coordinate_rows), default=0)
    extracted_width = max((len(row) for row in compact_pdf_table(extracted_rows)), default=0)
    if coordinate_width < 5 or coordinate_width > 24:
        return False
    if coordinate_width < extracted_width + 3:
        return False
    total_cells = max(coordinate_width * len(coordinate_rows), 1)
    filled_cells = sum(1 for row in coordinate_rows for cell in row if cell)
    fill_ratio = filled_cells / total_cells
    if fill_ratio < 0.18:
        return False
    if coordinate_width >= max(extracted_width * 4, 10) and pdf_table_numeric_cell_ratio(coordinate_rows) < 0.16:
        return False
    informative_rows = sum(1 for row in coordinate_rows if sum(1 for cell in row if cell) >= 3)
    return informative_rows >= 3


def should_use_grid_pdf_table(extracted_rows: list[list[str]], grid_rows: list[list[str]]) -> bool:
    if len(grid_rows) < 2:
        return False
    grid_width = max((len(row) for row in grid_rows), default=0)
    extracted_width = max((len(row) for row in compact_pdf_table(extracted_rows)), default=0)
    if grid_width < 2 or grid_width > 40:
        return False
    if grid_width < extracted_width and extracted_width <= 12:
        return False
    filled = sum(1 for row in grid_rows for cell in row if cell)
    if grid_table_has_stacked_numeric_cells(grid_rows):
        return False
    return filled >= 4


def grid_table_has_stacked_numeric_cells(rows: list[list[str]]) -> bool:
    filled = [cell for row in rows for cell in row if cell]
    if not filled:
        return False
    stacked = [cell for cell in filled if is_stacked_numeric_pdf_cell(cell)]
    if len(stacked) >= 3 and len(stacked) / len(filled) >= 0.12:
        return True
    for row in rows[1:]:
        row_stacked = sum(1 for cell in row if is_stacked_numeric_pdf_cell(cell))
        if row_stacked >= 3:
            return True
    return False


def is_stacked_numeric_pdf_cell(cell: str) -> bool:
    parts = [part.strip() for part in cell.split("/") if part.strip()]
    if len(parts) < 3:
        return False
    numeric_parts = sum(1 for part in parts if re.search(r"\d", part))
    return numeric_parts >= 3


def reconstruct_pdf_grid_table_from_page(page, bbox: tuple[float, float, float, float] | None) -> list[list[str]]:
    if page is None or bbox is None:
        return []
    return reconstruct_pdf_grid_table(
        word_items=pdf_word_items_from_page(page, bbox),
        line_segments=pdf_table_line_segments_from_page(page, bbox),
        bbox=bbox,
    )


def reconstruct_pdf_grid_table(
    word_items: list[tuple[float, float, float, float, str]],
    line_segments: list[PdfLineSegment],
    bbox: tuple[float, float, float, float] | None = None,
) -> list[list[str]]:
    if not word_items or not line_segments:
        return []
    vertical = [segment for segment in line_segments if segment.orientation == "v"]
    horizontal = [segment for segment in line_segments if segment.orientation == "h"]
    if len(vertical) < 2 or len(horizontal) < 2:
        return []
    x_edges = clustered_pdf_line_positions(vertical)
    y_edges = clustered_pdf_line_positions(horizontal)
    if bbox is not None:
        x_edges = add_missing_pdf_bbox_edges(x_edges, bbox[0], bbox[2])
        y_edges = add_missing_pdf_bbox_edges(y_edges, bbox[1], bbox[3])
    x_edges = sorted(unique_close_pdf_positions(x_edges))
    y_edges = sorted(unique_close_pdf_positions(y_edges))
    if len(x_edges) < 2 or len(y_edges) < 2:
        return []
    if not pdf_grid_edges_have_coverage(x_edges, y_edges, vertical, horizontal):
        return []
    rows: list[list[str]] = []
    for row_index in range(len(y_edges) - 1):
        row: list[str] = []
        y0, y1 = y_edges[row_index], y_edges[row_index + 1]
        for column_index in range(len(x_edges) - 1):
            x0, x1 = x_edges[column_index], x_edges[column_index + 1]
            text = text_in_pdf_grid_cell(word_items, x0, y0, x1, y1)
            row.append(text)
        if any(row):
            rows.append(row)
    return remove_duplicate_pdf_table_rows(rows)


def pdf_table_line_segments_from_page(page, bbox: tuple[float, float, float, float] | None) -> list[PdfLineSegment]:
    if page is None or bbox is None:
        return []
    try:
        drawings = page.get_drawings()
    except Exception:
        return []
    segments: list[PdfLineSegment] = []
    for drawing in drawings:
        for item in drawing.get("items", []):
            if not item or item[0] != "l" or len(item) < 3:
                continue
            p1, p2 = item[1], item[2]
            x0, y0 = float(p1.x), float(p1.y)
            x1, y1 = float(p2.x), float(p2.y)
            if not line_segment_overlaps_bbox(x0, y0, x1, y1, bbox):
                continue
            if abs(y0 - y1) <= 1.2 and abs(x1 - x0) >= 12:
                start, end = sorted((max(min(x0, x1), bbox[0]), min(max(x0, x1), bbox[2])))
                segments.append(PdfLineSegment("h", (y0 + y1) / 2, start, end))
            elif abs(x0 - x1) <= 1.2 and abs(y1 - y0) >= 12:
                start, end = sorted((max(min(y0, y1), bbox[1]), min(max(y0, y1), bbox[3])))
                segments.append(PdfLineSegment("v", (x0 + x1) / 2, start, end))
    return segments


def line_segment_overlaps_bbox(x0: float, y0: float, x1: float, y1: float, bbox: tuple[float, float, float, float]) -> bool:
    left, top, right, bottom = bbox
    return max(min(x0, x1), left) <= min(max(x0, x1), right) and max(min(y0, y1), top) <= min(max(y0, y1), bottom)


def clustered_pdf_line_positions(segments: list[PdfLineSegment]) -> list[float]:
    clusters: list[list[float]] = []
    for position in sorted(segment.position for segment in segments):
        if not clusters or abs(position - median(clusters[-1])) > 2.0:
            clusters.append([position])
        else:
            clusters[-1].append(position)
    return [median(cluster) for cluster in clusters]


def add_missing_pdf_bbox_edges(edges: list[float], start: float, end: float) -> list[float]:
    result = list(edges)
    if not result or min(abs(edge - start) for edge in result) > 3:
        result.append(start)
    if not result or min(abs(edge - end) for edge in result) > 3:
        result.append(end)
    return result


def unique_close_pdf_positions(values: list[float], tolerance: float = 2.0) -> list[float]:
    result: list[float] = []
    for value in sorted(values):
        if not result or abs(value - result[-1]) > tolerance:
            result.append(value)
        else:
            result[-1] = (result[-1] + value) / 2
    return result


def pdf_grid_edges_have_coverage(
    x_edges: list[float],
    y_edges: list[float],
    vertical: list[PdfLineSegment],
    horizontal: list[PdfLineSegment],
) -> bool:
    top, bottom = y_edges[0], y_edges[-1]
    left, right = x_edges[0], x_edges[-1]
    vertical_coverage = sum(1 for edge in x_edges if line_segments_cover_span(vertical, edge, top, bottom))
    horizontal_coverage = sum(1 for edge in y_edges if line_segments_cover_span(horizontal, edge, left, right))
    return vertical_coverage >= max(2, len(x_edges) // 2) and horizontal_coverage >= max(2, len(y_edges) // 2)


def line_segments_cover_span(segments: list[PdfLineSegment], position: float, start: float, end: float) -> bool:
    matching = [segment for segment in segments if abs(segment.position - position) <= 2.5]
    if not matching:
        return False
    covered = 0.0
    for segment in matching:
        covered += max(0.0, min(segment.end, end) - max(segment.start, start))
    return covered >= (end - start) * 0.55


def text_in_pdf_grid_cell(word_items: list[tuple[float, float, float, float, str]], x0: float, y0: float, x1: float, y1: float) -> str:
    words = []
    for word in word_items:
        center_x = (word[0] + word[2]) / 2
        center_y = (word[1] + word[3]) / 2
        if x0 - 1 <= center_x <= x1 + 1 and y0 - 1 <= center_y <= y1 + 1:
            words.append(word)
    if not words:
        return ""
    lines = group_pdf_word_items_into_lines(words)
    return " / ".join(normalize_text(" ".join(word[4] for word in line)) for line in lines if line)


def reconstruct_pdf_table_from_words(page, bbox: tuple[float, float, float, float] | None) -> list[list[str]]:
    return reconstruct_pdf_table_from_word_items(pdf_word_items_from_page(page, bbox))


def pdf_word_items_from_page(page, bbox: tuple[float, float, float, float] | None) -> list[tuple[float, float, float, float, str]]:
    if page is None or bbox is None:
        return []
    try:
        import fitz

        words = page.get_text("words", clip=fitz.Rect(*bbox))
    except Exception:
        return []
    return [
        (float(word[0]), float(word[1]), float(word[2]), float(word[3]), normalize_text(str(word[4])))
        for word in words
        if len(word) >= 5 and normalize_text(str(word[4]))
    ]


def reconstruct_pdf_table_from_word_items(word_items: list[tuple[float, float, float, float, str]]) -> list[list[str]]:
    lines = group_pdf_word_items_into_lines(word_items)
    if len(lines) < 3:
        return []
    anchors = infer_pdf_table_column_anchors(lines)
    if len(anchors) < 2:
        return []
    rows: list[list[str]] = []
    for line in lines:
        cells = [""] * len(anchors)
        for word in sorted(line, key=lambda item: item[0]):
            column = nearest_pdf_table_column(word[0], anchors)
            cells[column] = normalize_text(f"{cells[column]} {word[4]}")
        if any(cells):
            rows.append(cells)
    rows = attach_standalone_pdf_row_labels(rows)
    rows = merge_pdf_coordinate_header_rows(rows)
    return remove_sparse_pdf_coordinate_rows(rows)


def apply_extracted_row_labels_to_coordinate_rows(extracted_rows: list[list[str]], coordinate_rows: list[list[str]]) -> list[list[str]]:
    if len(coordinate_rows) < 2:
        return coordinate_rows
    expanded = expand_multiline_pdf_table_rows(compact_pdf_table(extracted_rows))
    if len(expanded) < 2:
        return coordinate_rows
    result = [coordinate_rows[0][:]] + [row[:] for row in coordinate_rows[1:]]
    labels = [row[0] for row in expanded[1:]]
    for index, label in enumerate(labels, start=1):
        if index >= len(result):
            break
        result[index][0] = label
    return result


def group_pdf_word_items_into_lines(
    word_items: list[tuple[float, float, float, float, str]],
) -> list[list[tuple[float, float, float, float, str]]]:
    if not word_items:
        return []
    heights = [item[3] - item[1] for item in word_items if item[3] > item[1]]
    tolerance = max(3.2, min(6.5, median(heights) * 0.45 if heights else 3.5))
    grouped: list[list[tuple[float, float, float, float, str]]] = []
    for item in sorted(word_items, key=lambda value: ((value[1] + value[3]) / 2, value[0])):
        center_y = (item[1] + item[3]) / 2
        if grouped:
            last_center = median([(word[1] + word[3]) / 2 for word in grouped[-1]])
            if abs(center_y - last_center) <= tolerance:
                grouped[-1].append(item)
                continue
        grouped.append([item])
    return [sorted(line, key=lambda value: value[0]) for line in grouped]


def infer_pdf_table_column_anchors(lines: list[list[tuple[float, float, float, float, str]]]) -> list[float]:
    x_values = sorted(word[0] for line in lines for word in line)
    if not x_values:
        return []
    char_widths = [
        (word[2] - word[0]) / max(len(word[4]), 1)
        for line in lines
        for word in line
        if word[2] > word[0]
    ]
    tolerance = max(4.5, min(7.5, (median(char_widths) if char_widths else 4.0) * 1.35))
    clusters: list[list[float]] = []
    for x in x_values:
        if not clusters or abs(x - median(clusters[-1])) > tolerance:
            clusters.append([x])
        else:
            clusters[-1].append(x)
    supported = [median(cluster) for cluster in clusters if len(cluster) >= 2]
    if len(supported) >= 2:
        return supported
    return [median(cluster) for cluster in clusters]


def nearest_pdf_table_column(x: float, anchors: list[float]) -> int:
    return min(range(len(anchors)), key=lambda index: abs(x - anchors[index]))


def attach_standalone_pdf_row_labels(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return rows
    result = [row[:] for row in rows]
    remove_indexes: set[int] = set()
    for index, row in enumerate(result):
        if index in remove_indexes or not row or not row[0]:
            continue
        if any(row[1:]):
            continue
        target = pdf_row_label_target(result, index)
        if target is None or target == index:
            continue
        if not result[target][0]:
            result[target][0] = row[0]
            remove_indexes.add(index)
    return [row for index, row in enumerate(result) if index not in remove_indexes]


def pdf_row_label_target(rows: list[list[str]], label_index: int) -> int | None:
    if label_index > 0 and not rows[label_index - 1][0] and any(rows[label_index - 1][1:]):
        return label_index - 1
    for index in range(label_index + 1, len(rows)):
        if not rows[index][0] and any(rows[index][1:]):
            return index
        if rows[index][0]:
            break
    return None


def merge_pdf_coordinate_header_rows(rows: list[list[str]]) -> list[list[str]]:
    if len(rows) < 4:
        return rows
    data_index = first_pdf_coordinate_data_row_index(rows)
    if data_index <= 1 or data_index > 4:
        return rows
    width = max(len(row) for row in rows)
    header = [""] * width
    for row in rows[:data_index]:
        padded = (row + [""] * (width - len(row)))[:width]
        for index, cell in enumerate(padded):
            if cell:
                header[index] = normalize_text(f"{header[index]} {cell}")
    data_rows = rows[data_index:]
    header = repair_pdf_table_header_fragments(header)
    header = align_pdf_header_to_data_columns(header, data_rows)
    return [header] + data_rows


def first_pdf_coordinate_data_row_index(rows: list[list[str]]) -> int:
    for index, row in enumerate(rows):
        filled = [cell for cell in row if cell]
        if len(filled) < 3:
            continue
        numeric = sum(1 for cell in filled if is_pdf_numeric_data_cell(cell))
        if numeric >= max(2, len(filled) // 2):
            return index
    return 0


def is_pdf_numeric_data_cell(cell: str) -> bool:
    text = normalize_text(cell)
    if not text:
        return False
    if re.fullmatch(r"[A-Za-z]+-?\d(?:\.\d+)?", text):
        return False
    return bool(re.search(r"(?:^|[~$€£±<>=\s-])\d[\d,.]*(?:%|K|M|B|x10\d*)?", text))


def remove_sparse_pdf_coordinate_rows(rows: list[list[str]]) -> list[list[str]]:
    result: list[list[str]] = []
    for index, row in enumerate(rows):
        filled = sum(1 for cell in row if cell)
        if filled == 0:
            continue
        if index > 0 and filled == 1 and len(row) > 8:
            continue
        result.append(row)
    return result


def expand_multiline_pdf_table_rows(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return rows
    width = max(len(row) for row in rows)
    padded = [(row + [""] * (width - len(row)))[:width] for row in rows]
    if should_merge_pdf_table_header_rows(padded):
        header = merge_pdf_table_header_rows(padded[0], padded[1])
        data_rows = padded[2:]
    else:
        header = flatten_pdf_header_row(padded[0])
        data_rows = padded[1:]
    expanded: list[list[str]] = [header]
    for row in data_rows:
        row_lines = [pdf_table_cell_lines(cell) for cell in row]
        subrow_count = pdf_table_multiline_subrow_count(row_lines)
        if subrow_count <= 1:
            expanded.append([flatten_pdf_table_cell(cell) for cell in row])
            continue
        for index in range(subrow_count):
            expanded_row: list[str] = []
            for lines in row_lines:
                if len(lines) == subrow_count:
                    expanded_row.append(lines[index])
                elif len(lines) > 1 and len(lines) >= subrow_count - 1:
                    expanded_row.append(lines[index] if index < len(lines) else "")
                elif len(lines) == 1:
                    expanded_row.append(lines[0] if index == 0 else "")
                else:
                    expanded_row.append(" / ".join(lines) if index == 0 else "")
            expanded.append(expanded_row)
    return remove_duplicate_pdf_table_rows(expanded)


def flatten_pdf_header_row(row: list[str]) -> list[str]:
    return [flatten_pdf_table_cell(cell, " / ") for cell in row]


def should_merge_pdf_table_header_rows(rows: list[list[str]]) -> bool:
    if len(rows) < 3:
        return False
    width = max((len(row) for row in rows), default=0)
    if width < 5:
        return False
    first_non_empty = sum(1 for cell in rows[0] if cell)
    second_non_empty = sum(1 for cell in rows[1] if cell)
    return first_non_empty <= max(2, width // 4) and second_non_empty >= max(3, width // 2)


def merge_pdf_table_header_rows(first: list[str], second: list[str]) -> list[str]:
    width = max(len(first), len(second))
    first = (first + [""] * (width - len(first)))[:width]
    second = (second + [""] * (width - len(second)))[:width]
    merged: list[str] = []
    for top, bottom in zip(first, second, strict=False):
        top_text = flatten_pdf_table_cell(top, " / ")
        bottom_text = flatten_pdf_table_cell(bottom)
        merged.append(bottom_text or top_text)
    return repair_pdf_table_header_fragments(merged)


def repair_pdf_table_header_fragments(header: list[str]) -> list[str]:
    result = header[:]
    for index, cell in enumerate(result):
        text = normalize_text(cell)
        if text.lower() != "(non-":
            continue
        target = None
        for candidate in range(index - 1, -1, -1):
            if "contaminated)" in result[candidate].lower():
                target = candidate
                break
            if result[candidate]:
                break
        if target is None:
            continue
        result[target] = re.sub(r"\s*contaminated\)", " (non-contaminated)", result[target], flags=re.I)
        result[index] = ""
    return result


def align_pdf_header_to_data_columns(header: list[str], data_rows: list[list[str]]) -> list[str]:
    if not data_rows:
        return header
    result = header[:]
    width = len(result)
    data_counts = []
    for index in range(width):
        data_counts.append(sum(1 for row in data_rows if index < len(row) and row[index]))
    for index, cell in enumerate(header[:-1]):
        if not cell or data_counts[index] > 0:
            continue
        for target in range(index + 1, min(width, index + 3)):
            if result[target] or data_counts[target] == 0:
                continue
            result[target] = cell
            result[index] = ""
            break
    return result


def table_has_expandable_multiline_rows(rows: list[list[str]]) -> bool:
    for row in rows[1:]:
        row_lines = [pdf_table_cell_lines(cell) for cell in row]
        if pdf_table_multiline_subrow_count(row_lines) > 1:
            return True
    return False


def pdf_table_multiline_subrow_count(row_lines: list[list[str]]) -> int:
    counts = [len(lines) for lines in row_lines if len(lines) > 1]
    if not counts:
        return 1
    best = max(set(counts), key=lambda value: (counts.count(value), value))
    if len(counts) == 2 and min(counts) >= 4 and abs(counts[0] - counts[1]) <= 1:
        return best
    if counts.count(best) >= 2:
        return best
    return best if len(counts) >= 3 else 1


def remove_duplicate_pdf_table_rows(rows: list[list[str]]) -> list[list[str]]:
    result: list[list[str]] = []
    seen: set[tuple[str, ...]] = set()
    for row in rows:
        if not any(row):
            continue
        key = tuple(row)
        if key in seen:
            continue
        seen.add(key)
        result.append(row)
    return result


def pdf_table_layout_lines_from_page(page, bbox: tuple[float, float, float, float] | None) -> list[str]:
    if page is None or bbox is None:
        return []
    try:
        import fitz

        words = page.get_text("words", clip=fitz.Rect(*bbox))
    except Exception:
        return []
    if not words:
        return []
    word_items = [
        (float(word[0]), float(word[1]), float(word[2]), float(word[3]), normalize_text(str(word[4])))
        for word in words
        if len(word) >= 5 and normalize_text(str(word[4]))
    ]
    if not word_items:
        return []
    char_widths = [(x1 - x0) / max(len(text), 1) for x0, _y0, x1, _y1, text in word_items if x1 > x0]
    unit = max(2.8, min(8.0, median(char_widths) if char_widths else 4.0))
    grouped: list[list[tuple[float, float, float, float, str]]] = []
    for item in sorted(word_items, key=lambda value: ((value[1] + value[3]) / 2, value[0])):
        center_y = (item[1] + item[3]) / 2
        if grouped:
            last_center = median([(word[1] + word[3]) / 2 for word in grouped[-1]])
            if abs(center_y - last_center) <= 3.5:
                grouped[-1].append(item)
                continue
        grouped.append([item])
    x_origin = min(item[0] for item in word_items)
    lines: list[str] = []
    for group in grouped:
        pieces: list[str] = []
        current_col = 0
        for x0, _y0, _x1, _y1, text in sorted(group, key=lambda value: value[0]):
            col = max(0, int(round((x0 - x_origin) / unit)))
            gap = max(1, col - current_col)
            pieces.append(" " * gap + text)
            current_col = col + len(text)
        line = "".join(pieces).rstrip()
        if line:
            lines.append(line)
    return lines


def bbox_overlaps_any(bbox: tuple | list, others: list[tuple[float, float, float, float]]) -> bool:
    current = tuple(float(value) for value in bbox)
    area = max(0.0, current[2] - current[0]) * max(0.0, current[3] - current[1])
    if area <= 0:
        return False
    for other in others:
        x0 = max(current[0], other[0])
        y0 = max(current[1], other[1])
        x1 = min(current[2], other[2])
        y1 = min(current[3], other[3])
        overlap = max(0.0, x1 - x0) * max(0.0, y1 - y0)
        if overlap / area >= 0.35:
            return True
    return False


def pdf_image_block(
    block: dict,
    page_number: int,
    image_index: int,
    assets_dir: Path | None,
    used_image_names: set[str] | None,
) -> PdfTextBlock | None:
    image_bytes = block.get("image")
    if not assets_dir or not image_bytes:
        return None
    width = int(block.get("width") or 0)
    height = int(block.get("height") or 0)
    if width <= 8 or height <= 8:
        return None
    extension = str(block.get("ext") or "png").lower().lstrip(".")
    if extension not in {"png", "jpg", "jpeg", "webp", "gif"}:
        extension = "png"
    used = used_image_names if used_image_names is not None else set()
    filename = unique_filename(f"pdf-page-{page_number}-image-{image_index}.{extension}", used)
    destination = assets_dir / filename
    destination.write_bytes(image_bytes)
    relative_path = f"./assets/{filename}"
    alt = Path(filename).stem
    bbox = tuple(float(value) for value in block.get("bbox", (0, 0, 0, 0)))
    return PdfTextBlock(
        page=page_number,
        text=f"![{alt}]({relative_path})",
        bbox=bbox,
        font_size=0,
        is_image=True,
        image_path=relative_path,
    )


def join_pdf_spans(spans: list[dict]) -> str:
    if not spans:
        return ""
    pieces: list[str] = []
    previous_x1: float | None = None
    previous_size = float(spans[0].get("size") or 10)
    for span in spans:
        text = str(span.get("text") or "")
        if not text:
            continue
        bbox = span.get("bbox", (0, 0, 0, 0))
        x0 = float(bbox[0])
        if previous_x1 is not None and not text.startswith((" ", "\t")):
            gap = x0 - previous_x1
            if gap > max(previous_size * 0.18, 1.2):
                pieces.append(" ")
        pieces.append(text)
        previous_x1 = float(bbox[2])
        previous_size = float(span.get("size") or previous_size)
    return "".join(pieces)


def is_pdf_math_font(font: str) -> bool:
    font = font.lower()
    return any(marker in font for marker in ["cmmi", "cmsy", "msbm", "symbol", "math"])


def is_pdf_formula_block(text: str, math_ratio: float) -> bool:
    compact = normalize_text(text.replace("\n", " "))
    if not compact:
        return False
    if is_pdf_formula_explanation(compact):
        return False
    if math_ratio >= 0.45 and len(compact) <= 220 and formula_symbol_density(compact) >= 0.08:
        return True
    if math_ratio >= 0.55 and len(compact) <= 160:
        return True
    if re.search(r"[=∈⊤→∑√∫≤≥≪×−]", compact) and re.search(r"\(\d+\)\s*$", compact) and len(compact) <= 260:
        return True
    return False


def is_standalone_pdf_page_number(line: str, page_number: int) -> bool:
    if not line.isdigit():
        return False
    value = int(line)
    return value == page_number or 1 <= value <= 30


def is_pdf_formula_explanation(text: str) -> bool:
    normalized = normalize_text(text).lower()
    if normalized.startswith(("where ", "when ", "for ")):
        return True
    words = re.findall(r"[a-z]{2,}", normalized)
    if len(words) < 6:
        return False
    if re.search(r"\b(is|are|denotes|represents)\s+(?:the|a|an)\b", normalized) and formula_symbol_density(text) < 0.22:
        return True
    explanation_markers = {"is", "are", "the", "and", "of", "for", "with", "from", "to"}
    marker_count = sum(1 for word in words if word in explanation_markers)
    return marker_count >= 3 and formula_symbol_density(text) < 0.18


def recurring_pdf_noise_keys(page_blocks: list[list[PdfTextBlock]]) -> set[str]:
    counts: dict[str, int] = {}
    for blocks in page_blocks:
        seen_on_page: set[str] = set()
        for block in blocks:
            key = pdf_block_key(block.text)
            if key:
                seen_on_page.add(key)
        for key in seen_on_page:
            counts[key] = counts.get(key, 0) + 1
    page_count = len(page_blocks)
    threshold = max(3, int(page_count * 0.45))
    return {key for key, count in counts.items() if count >= threshold}


def pdf_block_key(text: str) -> str:
    text = normalize_text(text)
    if len(text) > 140:
        return ""
    return comparable_pdf_text(text)


def should_drop_pdf_block(block: PdfTextBlock, page_height: float, recurring: set[str]) -> bool:
    text = normalize_text(block.text)
    key = pdf_block_key(text)
    y0, y1 = block.bbox[1], block.bbox[3]
    in_margin = y0 < page_height * 0.075 or y1 > page_height * 0.94
    if key and key in recurring and in_margin:
        return True
    if is_pdf_noise_line(text):
        return True
    if re.fullmatch(r"\d{1,4}", text) and in_margin:
        return True
    return False


def strip_pdf_title_blocks(blocks: list[PdfTextBlock], title: str) -> list[PdfTextBlock]:
    if not blocks or not title.strip():
        return blocks
    title_key = comparable_pdf_text(title)
    accumulated = ""
    title_end = -1
    for index, block in enumerate(blocks[:10]):
        if block.is_image:
            continue
        block_key = comparable_pdf_text(block.text)
        if title_key and (block_key == title_key or title_key in block_key):
            title_end = index
            break
        accumulated = normalize_text(f"{accumulated} {block.text.replace(chr(10), ' ')}")
        if title_key and comparable_pdf_text(accumulated).startswith(title_key[: max(24, min(len(title_key), 80))]):
            title_end = index
            break
    if title_end < 0:
        return blocks
    for index in range(title_end + 1, min(len(blocks), title_end + 12)):
        if normalize_text(blocks[index].text).lower() == "abstract":
            return blocks[index:]
    return blocks[title_end + 1 :]


def render_pdf_blocks(blocks: list[PdfTextBlock], body_size: float) -> str:
    parts: list[str] = []
    paragraph: list[str] = []
    pending_formula: str | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            parts.append(normalize_text(" ".join(paragraph)))
            paragraph = []

    def flush_formula() -> None:
        nonlocal pending_formula
        if pending_formula:
            parts.append(render_pdf_formula(pending_formula))
            pending_formula = None

    def append_formula_line(line: str) -> None:
        nonlocal pending_formula
        if pending_formula:
            separator = " " if is_formula_number_line(line) else "\n"
            pending_formula = f"{pending_formula}{separator}{line}"
        else:
            pending_formula = line

    for block in blocks:
        if block.is_image or block.is_table:
            flush_paragraph()
            flush_formula()
            if block.text:
                parts.append(block.text)
            continue
        for raw_line in block.text.splitlines():
            line = normalize_text(raw_line)
            if not line:
                flush_paragraph()
                flush_formula()
                continue
            line = repair_pdf_line_artifacts(line)
            if is_pdf_noise_line(line):
                flush_paragraph()
                flush_formula()
                continue
            if block.is_formula:
                flush_paragraph()
                append_formula_line(line)
                if not is_formula_number_line(line) and not may_formula_continue(line):
                    flush_formula()
                continue
            if is_standalone_pdf_page_number(line, block.page):
                flush_paragraph()
                flush_formula()
                continue
            if pending_formula and is_pdf_formula_tail(line):
                append_formula_line(line)
                continue
            if is_pdf_formula_line(line):
                flush_paragraph()
                append_formula_line(line)
                if not is_formula_number_line(line) and not may_formula_continue(line):
                    flush_formula()
                continue
            if is_pdf_caption_line(line):
                flush_paragraph()
                flush_formula()
                parts.append(f"> {line}")
                continue
            if is_pdf_heading_line(line) or is_layout_pdf_heading(line, block, body_size):
                flush_paragraph()
                flush_formula()
                parts.append(f"### {line}")
                continue
            if is_list_line(line):
                flush_paragraph()
                flush_formula()
                parts.append(normalize_pdf_list_line(line))
                continue
            flush_formula()
            paragraph.append(line)
            if should_end_pdf_paragraph(line):
                flush_paragraph()
    flush_paragraph()
    flush_formula()
    return clean_markdown("\n\n".join(parts), "")


def export_pdf_images(document, assets_dir: Path) -> list[str]:
    resources: list[str] = []
    used: set[str] = set()
    for page_index in range(document.page_count):
        page = document[page_index]
        for image_index, image_info in enumerate(page.get_images(full=True), start=1):
            xref = image_info[0]
            try:
                image = document.extract_image(xref)
            except Exception:
                continue
            image_bytes = image.get("image")
            if not image_bytes:
                continue
            extension = image.get("ext") or "png"
            width = int(image.get("width") or 0)
            height = int(image.get("height") or 0)
            if width <= 8 or height <= 8:
                continue
            filename = unique_filename(f"pdf-page-{page_index + 1}-image-{image_index}.{extension}", used)
            destination = assets_dir / filename
            destination.write_bytes(image_bytes)
            resources.append(f"./assets/{filename}")
    return resources


def extract_pdf_with_pypdf(reader: PdfReader, info: PdfDocumentInfo) -> PdfExtraction:
    parts: list[str] = []
    stats = DocumentStats(pages=len(reader.pages))
    for index, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        if index == 1:
            raw_text = strip_pdf_title_block(raw_text, info.title)
        text = normalize_pdf_text(raw_text)
        if not text:
            continue
        stats.characters += len(text)
        stats.paragraphs += count_paragraphs(text)
        stats.headings += count_markdown_headings(text)
        stats.formulas += count_markdown_formulas(text)
        parts.append(f"## Page {index}\n\n{text}")
    body = repair_pdf_text_artifacts(repair_pdf_page_continuations("\n\n".join(parts)))
    metadata = {"pdf_engine": "pypdf"}
    if not body:
        metadata["warning"] = "No extractable text found. Scanned PDFs need OCR, which is planned for a later phase."
    return PdfExtraction(body=body, resources=[], stats=stats, metadata=metadata)


def median(values: list[float]) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    midpoint = len(ordered) // 2
    if len(ordered) % 2:
        return ordered[midpoint]
    return (ordered[midpoint - 1] + ordered[midpoint]) / 2


def pdf_document_info(reader: PdfReader, path: Path) -> PdfDocumentInfo:
    metadata = getattr(reader, "metadata", None) or {}
    title = pdf_metadata_value(metadata.get("/Title")) or path.stem
    author = pdf_metadata_value(metadata.get("/Author"))
    source_url = pdf_metadata_value(metadata.get("/DOI")) or pdf_metadata_value(metadata.get("/arXivID"))
    created_at = parse_pdf_date(pdf_metadata_value(metadata.get("/CreationDate")))
    return PdfDocumentInfo(title=title, author=author, source_url=source_url, created_at=created_at)


def infer_pdf_title_from_blocks(blocks: list[PdfTextBlock], fallback: str, body_size: float) -> str:
    candidates: list[tuple[float, str]] = []
    for block in blocks[:20]:
        if block.is_image:
            continue
        y0 = block.bbox[1]
        if y0 > 420 and candidates:
            break
        text = normalize_text(block.text.replace("\n", " "))
        if not is_plausible_pdf_title(text):
            continue
        font_delta = block.font_size - body_size
        if font_delta < 1.4:
            continue
        score = font_delta * 25 - y0 / 18 + min(len(text), 160) / 20
        if y0 < 80:
            score -= 35
        candidates.append((score, text))
    if not candidates:
        return fallback
    return max(candidates, key=lambda item: item[0])[1]


def is_plausible_pdf_title(text: str) -> bool:
    if not 8 <= len(text) <= 220:
        return False
    lowered = text.lower()
    if lowered.startswith(("arxiv:", "provided proper attribution", "copyright", "abstract", "contents")):
        return False
    if "@" in text or re.search(r"https?://", text):
        return False
    if re.fullmatch(r"[\d\s.:-]+", text):
        return False
    return True


def should_replace_pdf_title(current: str, path: Path, inferred: str) -> bool:
    if not inferred.strip():
        return False
    current_key = comparable_pdf_text(current)
    inferred_key = comparable_pdf_text(inferred)
    if not inferred_key:
        return False
    path_key = comparable_pdf_text(path.stem)
    if current_key == path_key:
        return normalize_text(current) != normalize_text(inferred)
    if inferred_key == current_key:
        return False
    return current.lower().endswith(".pdf")


def pdf_metadata_value(value) -> str | None:
    if value is None:
        return None
    text = normalize_text(str(value))
    return text or None


def parse_pdf_date(value: str | None) -> str | None:
    if not value:
        return None
    match = re.match(r"D:(\d{4})(\d{2})(\d{2})", value)
    if match:
        return "-".join(match.groups())
    return value


def repair_pdf_page_continuations(markdown: str) -> str:
    pattern = re.compile(r"([A-Za-z]{3,})-\n\n## Page (\d+)\n\n([a-z]{2,})")
    return pattern.sub(lambda match: f"{match.group(1)}{match.group(3)}\n\n## Page {match.group(2)}\n\n", markdown)


def repair_pdf_text_artifacts(text: str) -> str:
    text = re.sub(r"\b([A-Za-z]+)-\s+([a-z]{2,})\b", r"\1\2", text)
    text = re.sub(r"\b(\d+)\.([A-Z])", r"\1. \2", text)
    text = re.sub(r"(?<=\S)(Figure\s+\d+\.)", r"\n\n> \1", text)
    text = re.sub(r"(?<=\S)(Table\s+\d+\.)", r"\n\n> \1", text)
    text = merge_adjacent_pdf_formula_blocks(text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def merge_adjacent_pdf_formula_blocks(text: str) -> str:
    pattern = re.compile(r"\$\$\n(?P<left>.+?)\n\$\$\n\n\$\$\n(?P<right>.+?)\n\$\$", re.S)
    previous = None
    while previous != text:
        previous = text

        def replace(match: re.Match) -> str:
            left = match.group("left").strip()
            right = match.group("right").strip()
            if should_merge_formula_parts(left, right):
                return "$$\n" + left + "\n" + right + "\n$$"
            return match.group(0)

        text = pattern.sub(replace, text)
    return text


def should_merge_formula_parts(left: str, right: str) -> bool:
    if is_formula_number_line(right):
        return True
    if is_formula_number_line(left):
        return False
    if len(left) > 400 or len(right) > 400:
        return False
    return formula_symbol_density(left) >= 0.06 or formula_symbol_density(right) >= 0.06


def strip_pdf_title_block(text: str, title: str) -> str:
    if not text.strip() or not title.strip():
        return text
    lines = text.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    title_key = comparable_pdf_text(title)
    accumulated = ""
    title_end = -1
    for index, line in enumerate(lines[:12]):
        accumulated = normalize_text(f"{accumulated} {line}")
        if title_key and comparable_pdf_text(accumulated).startswith(title_key[: max(24, min(len(title_key), 80))]):
            title_end = index
            break
    if title_end < 0:
        return text

    start = title_end + 1
    for index in range(start, min(len(lines), start + 24)):
        if normalize_text(lines[index]).lower() == "abstract":
            return "\n".join(lines[index:])
    return "\n".join(lines[start:])


def comparable_pdf_text(value: str) -> str:
    return re.sub(r"[\W_]+", "", value, flags=re.UNICODE).lower()


def normalize_pdf_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(filter_pdf_noise_lines(text.splitlines()))
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    lines = [line.strip() for line in text.splitlines()]
    parts: list[str] = []
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            parts.append(normalize_text(" ".join(paragraph)))
            paragraph = []

    for raw_line in lines:
        line = normalize_text(raw_line)
        if not line:
            flush_paragraph()
            continue
        if is_pdf_noise_line(line):
            flush_paragraph()
            continue
        if is_pdf_heading_line(line) or is_legacy_numbered_pdf_heading(line):
            flush_paragraph()
            parts.append(f"### {line}")
            continue
        if is_list_line(line):
            flush_paragraph()
            parts.append(line)
            continue
        paragraph.append(line)
        if re.search(r"[。！？.!?]$", line):
            flush_paragraph()
    flush_paragraph()
    return clean_markdown("\n\n".join(parts), "")


def is_pdf_heading_line(line: str) -> bool:
    if not 3 <= len(line) <= 120:
        return False
    if is_bad_pdf_heading_candidate(line):
        return False
    if re.search(r"[。！？.!?;,，；]$", line):
        return False
    if re.search(r"[=∈⊤→▼ˆ¯]|^\W+$", line):
        return False
    if is_bullet_line(line):
        return False
    if re.match(r"^\d+(?:\.\d+)+\.?\s+[A-Z][\w–—:-]+(?:\s+\S+){0,10}$", line) and len(line.split()) <= 12:
        return True
    if re.match(r"^(?:[一二三四五六七八九十]+[、.．]|第[一二三四五六七八九十\d]+[章节])\s*\S+", line):
        return True
    if re.match(r"^(abstract|introduction|background|methods?|results?|discussion|conclusion|references|appendix)\b", line, re.I):
        return True
    letters = re.sub(r"[^A-Za-z]", "", line)
    if len(letters) >= 6 and letters.upper() == letters and not re.search(r"[\d,;:]", line):
        return True
    return False


def is_legacy_numbered_pdf_heading(line: str) -> bool:
    match = re.fullmatch(r"(\d{1,2})\.\s+([A-Z][\w–—-]+(?:\s+\S+){0,8})", line)
    if not match:
        return False
    title = match.group(2)
    if title.endswith("-"):
        return False
    if ":" in title or re.search(r"[,;，；]", title):
        return False
    if len(title.split()) > 8:
        return False
    return not looks_like_pdf_sentence(title)


def is_layout_pdf_heading(line: str, block: PdfTextBlock, body_size: float) -> bool:
    if not 3 <= len(line) <= 140:
        return False
    if is_bad_pdf_heading_candidate(line):
        return False
    if block.page == 1 and looks_like_pdf_cover_metadata(line):
        return False
    if re.search(r"[。！？.!?;,，；]$", line):
        return False
    if is_bullet_line(line) or is_pdf_caption_line(line):
        return False
    word_count = len(line.split())
    if re.match(r"^\d+[.)]\s+", line) and not (
        is_legacy_numbered_pdf_heading(line) or re.match(r"^\d+(?:\.\d+)+\.?\s+", line)
    ):
        return False
    if is_legacy_numbered_pdf_heading(line) and block.font_size >= body_size + 0.8 and word_count <= 8:
        return True
    if block.font_size >= body_size + 1.4 and word_count <= 12 and not looks_like_pdf_sentence(line):
        return True
    if block.is_bold and block.font_size >= body_size + 0.2 and word_count <= 10 and not looks_like_pdf_sentence(line):
        return True
    return False


def is_bad_pdf_heading_candidate(line: str) -> bool:
    if re.match(r"^\d{3,}\.", line):
        return True
    if is_pdf_formula_explanation(line):
        return True
    if ":" in line and len(line.split()) > 8:
        return True
    if len(line.split()) > 12 and looks_like_pdf_sentence(line):
        return True
    return False


def looks_like_pdf_sentence(line: str) -> bool:
    normalized = normalize_text(line).lower()
    words = re.findall(r"[a-z]{2,}", normalized)
    if len(words) < 6:
        return False
    sentence_markers = {"and", "or", "the", "that", "which", "with", "from", "those", "this", "these", "their"}
    marker_count = sum(1 for word in words if word in sentence_markers)
    return marker_count >= 2


def looks_like_pdf_cover_metadata(line: str) -> bool:
    text = normalize_text(line)
    if text.endswith("Authors:"):
        return True
    if re.search(r"^(URL/DOI|DOI|Author|Authors|Prepared by)$", text, re.I):
        return True
    if not re.fullmatch(r"[A-Z][A-Za-z.'-]+(?:\s+[A-Z]\.)?(?:\s+[A-Z][A-Za-z.'-]+){1,3}", text):
        return False
    words = text.split()
    return 3 <= len(words) <= 5 or any(word.endswith(".") and len(word) == 2 for word in words)


def is_pdf_caption_line(line: str) -> bool:
    return bool(re.match(r"^(Figure|Fig\.|Table)\s+\d+[\.:]\s+", line, re.I))


def is_pdf_formula_line(line: str) -> bool:
    if len(line) > 260:
        return False
    if is_pdf_formula_explanation(line):
        return False
    if is_formula_number_line(line):
        return True
    if re.search(r"[=∈⊤→∑√∫≤≥≪×−]", line) and (re.search(r"\(\d+\)\s*$", line) or formula_symbol_density(line) >= 0.08):
        return True
    return False


def is_pdf_formula_tail(line: str) -> bool:
    if len(line) > 100 or is_pdf_formula_explanation(line):
        return False
    if normalize_text(line).lower().startswith(("where ", "when ", "for ")):
        return False
    if re.fullmatch(r"[A-Za-z]{1,8}\([^)]{1,40}\)\.?", line):
        return True
    if formula_symbol_density(line) >= 0.16 and not looks_like_pdf_sentence(line):
        return True
    return False


def formula_symbol_density(line: str) -> float:
    compact = re.sub(r"\s+", "", line)
    if not compact:
        return 0.0
    symbols = len(re.findall(r"[=+\-*/^_∈⊤→∑√∫≤≥≪×−()]", compact))
    return symbols / len(compact)


def is_formula_number_line(line: str) -> bool:
    return bool(re.fullmatch(r"\(\d+[a-z]?\)", line.strip()))


def may_formula_continue(line: str) -> bool:
    return not re.search(r"\(\d+[a-z]?\)\s*$", line.strip())


def render_pdf_formula(line: str) -> str:
    return "$$\n" + line.strip() + "\n$$"


def normalize_pdf_list_line(line: str) -> str:
    match = re.match(r"^(\d+[.)])\s*(.+)$", line)
    if match:
        return f"{match.group(1)} {match.group(2)}"
    match = re.match(r"^([-*+•])\s*(.+)$", line)
    if match:
        return f"- {match.group(2)}"
    return line


def should_end_pdf_paragraph(line: str) -> bool:
    if re.search(r"[。！？.!?:：]$", line):
        return True
    if re.search(r"\)\s*$", line) and len(line) > 80:
        return True
    return False


def repair_pdf_line_artifacts(line: str) -> str:
    line = re.sub(r"\b(\d+)\.([A-Z])", r"\1. \2", line)
    line = re.sub(r"\b(Sec|Fig|Eq)\.\s*(\d)", r"\1. \2", line)
    return line


def is_list_line(line: str) -> bool:
    return bool(re.match(r"^([-*+•]|\d+[.)]|[a-zA-Z][.)])\s+", line))


def is_bullet_line(line: str) -> bool:
    return bool(re.match(r"^([-*+•]|[a-zA-Z][.)])\s+", line))


def filter_pdf_noise_lines(lines: list[str]) -> list[str]:
    filtered: list[str] = []
    total = len(lines)
    for index, raw_line in enumerate(lines):
        line = normalize_text(raw_line)
        if is_pdf_noise_line(line, index, total):
            continue
        filtered.append(raw_line)
    return filtered


def is_pdf_noise_line(line: str, index: int | None = None, total: int | None = None) -> bool:
    if re.match(r"^\d+\s+arxiv:\d{4}\.\d{4,5}v\d+\s+\[[^\]]+]\s+\d{1,2}\s+\w+\s+\d{4}$", line, re.I):
        return True
    if re.match(r"^arxiv:\d{4}\.\d{4,5}v\d+\s+\[[^\]]+]\s+\d{1,2}\s+\w+\s+\d{4}$", line, re.I):
        return True
    if index is not None and total is not None and re.fullmatch(r"\d{1,4}", line):
        return index <= 2 or index >= max(total - 3, 0)
    return False


def iter_docx_blocks(document: DocxDocument):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def paragraph_markdown(paragraph: Paragraph) -> str:
    text = paragraph_inline_markdown(paragraph)
    if not text:
        return ""
    style = (paragraph.style.name if paragraph.style else "").lower()
    level = heading_level(style)
    if level:
        return f"{'#' * level} {text}"
    return text


def heading_level(style: str) -> int | None:
    match = re.search(r"heading\s+([1-6])", style)
    if match:
        return int(match.group(1))
    match = re.search(r"标题\s*([1-6])", style)
    if match:
        return int(match.group(1))
    if "title" == style:
        return 1
    return None


def paragraph_inline_markdown(paragraph: Paragraph) -> str:
    chunks: list[str] = []
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            chunks.append(run_element_to_markdown(child))
        elif child.tag == qn("w:hyperlink"):
            chunks.append(hyperlink_element_to_markdown(paragraph, child))
    return normalize_text("".join(chunks))


def hyperlink_element_to_markdown(paragraph: Paragraph, element) -> str:
    rel_id = element.attrib.get(qn("r:id"))
    text = normalize_text("".join(run_element_to_text(run) for run in element if run.tag == qn("w:r")))
    if not text:
        return ""
    if not rel_id or rel_id not in paragraph.part.rels:
        return text
    target = paragraph.part.rels[rel_id].target_ref
    return f"[{text}]({target})"


def run_element_to_markdown(element) -> str:
    text = run_element_to_text(element)
    if not text:
        return ""
    rpr = element.find(qn("w:rPr"))
    bold = has_run_property(rpr, "w:b")
    italic = has_run_property(rpr, "w:i")
    code = run_has_style(rpr, {"code", "monospace", "source"})
    stripped = text.strip()
    if not stripped:
        return text
    prefix = text[: len(text) - len(text.lstrip())]
    suffix = text[len(text.rstrip()) :]
    content = stripped.replace("`", "\\`") if code else stripped
    if code:
        content = f"`{content}`"
    if bold and italic:
        content = f"***{content}***"
    elif bold:
        content = f"**{content}**"
    elif italic:
        content = f"*{content}*"
    return prefix + content + suffix


def run_element_to_text(element) -> str:
    parts: list[str] = []
    for child in element.iter():
        if child.tag == qn("w:t") and child.text:
            parts.append(child.text)
        elif child.tag == qn("w:tab"):
            parts.append("\t")
        elif child.tag == qn("w:br"):
            parts.append("\n")
    return "".join(parts)


def has_run_property(rpr, property_name: str) -> bool:
    if rpr is None:
        return False
    node = rpr.find(qn(property_name))
    if node is None:
        return False
    return node.attrib.get(qn("w:val"), "1") not in {"0", "false", "False"}


def run_has_style(rpr, markers: set[str]) -> bool:
    if rpr is None:
        return False
    style = rpr.find(qn("w:rStyle"))
    if style is None:
        return False
    value = style.attrib.get(qn("w:val"), "").lower()
    return any(marker in value for marker in markers)


def table_rows(table: Table) -> list[list[str]]:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not rows:
        return []
    width = max(len(row) for row in rows)
    return [(row + [""] * (width - len(row)))[:width] for row in rows]


def extract_docx_images(document: DocxDocument, assets_dir: Path) -> dict[str, str]:
    images: dict[str, str] = {}
    used_names: set[str] = set()
    for rel_id, rel in document.part.rels.items():
        if "image" not in rel.reltype:
            continue
        image_part = rel.target_part
        original_name = Path(image_part.partname).name
        filename = unique_filename(original_name, used_names)
        destination = assets_dir / filename
        destination.write_bytes(image_part.blob)
        images[rel_id] = f"./assets/{filename}"
    return images


def unique_filename(filename: str, used: set[str]) -> str:
    stem = Path(filename).stem or "image"
    suffix = Path(filename).suffix or ".bin"
    candidate = f"{stem}{suffix}"
    index = 1
    while candidate in used:
        index += 1
        candidate = f"{stem}-{index}{suffix}"
    used.add(candidate)
    return candidate


def paragraph_image_relationship_ids(paragraph: Paragraph) -> list[str]:
    ids: list[str] = []
    try:
        root = ElementTree.fromstring(paragraph._p.xml)
    except Exception:
        return ids
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    for blip in root.findall(".//a:blip", namespaces):
        rel_id = blip.attrib.get(f"{{{namespaces['r']}}}embed")
        if rel_id:
            ids.append(rel_id)
    return ids


def count_paragraphs(markdown: str) -> int:
    return len([part for part in re.split(r"\n\s*\n", markdown.strip()) if normalize_text(part)])


def count_markdown_headings(markdown: str) -> int:
    return len(re.findall(r"(?m)^#{1,6}\s+", markdown))


def count_markdown_tables(markdown: str) -> int:
    return len(re.findall(r"(?m)^\|.+\|\n\|(?:\s*:?-+:?\s*\|)+", markdown))


def count_markdown_table_rows(markdown: str) -> int:
    rows = 0
    lines = markdown.splitlines()
    for index, line in enumerate(lines):
        if not line.startswith("|"):
            continue
        if index > 0 and re.fullmatch(r"\|(?:\s*:?-{3,}:?\s*\|)+\s*", line):
            body_rows = 0
            for next_line in lines[index + 1 :]:
                if not next_line.startswith("|"):
                    break
                body_rows += 1
            rows += body_rows
    return rows


def count_markdown_formulas(markdown: str) -> int:
    return len(re.findall(r"(?s)\$\$\n.+?\n\$\$", markdown))


def stats_metadata(converter: str, stats: DocumentStats) -> dict:
    return {
        "converter": converter,
        "paragraph_count": stats.paragraphs,
        "heading_count": stats.headings,
        "table_count": stats.tables,
        "row_count": stats.rows,
        "image_count": stats.images,
        "formula_count": stats.formulas,
        "page_count": stats.pages,
        "character_count": stats.characters,
    }
