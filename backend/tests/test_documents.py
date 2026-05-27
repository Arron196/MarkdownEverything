from pathlib import Path

from docx import Document

from app.markdown import render_document_markdown
from app.converters.documents import (
    PdfLineSegment,
    PdfTextBlock,
    apply_extracted_row_labels_to_coordinate_rows,
    collapse_pdf_spanner_columns,
    convert_csv,
    convert_html,
    convert_pdf,
    convert_text,
    is_layout_pdf_heading,
    is_false_positive_pdf_paragraph_table,
    normalize_pdf_text,
    reconstruct_pdf_grid_table,
    reconstruct_pdf_table_from_word_items,
    render_coordinate_pdf_table,
    render_pdf_blocks,
    render_pdf_form_table,
    render_pdf_table_layout_block,
    render_pdf_markdown_table,
    should_prefer_extracted_pdf_table,
    should_render_pdf_table_as_layout_block,
    should_render_pdf_table_as_form,
    should_use_coordinate_pdf_table,
)
from app.converters.documents import convert_docx


def test_convert_text(tmp_path: Path):
    path = tmp_path / "note.txt"
    path.write_text("Hello Title\n\nWorld", encoding="utf-8")
    result = convert_text(path)
    assert result.title == "Hello Title"
    assert result.body == "World"
    assert result.metadata["converter"] == "text"


def test_convert_csv(tmp_path: Path):
    path = tmp_path / "data.csv"
    path.write_text("a,b\n1|x,2\n", encoding="utf-8")
    result = convert_csv(path)
    assert "| a | b |" in result.body
    assert "| 1\\|x | 2 |" in result.body
    assert result.metadata["table_count"] == 1


def test_convert_html_cleans_shell_noise_and_keeps_structure(tmp_path: Path):
    path = tmp_path / "page.html"
    path.write_text(
        """
        <html>
          <head><title>HTML Title</title></head>
          <body>
            <nav>Home Search Login</nav>
            <main><h1>Article Title</h1><p>Useful <strong>body</strong>.</p></main>
            <footer>Subscribe</footer>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = convert_html(path)

    assert result.title == "Article Title"
    rendered = render_document_markdown(title=result.title, source_type=result.source_type, body=result.body, summary="")
    assert rendered.count("# Article Title") == 1
    assert "**body**" in result.body
    assert "Home Search Login" not in result.body
    assert "Subscribe" not in result.body


def test_convert_html_prefers_main_content_container(tmp_path: Path):
    path = tmp_path / "page.html"
    path.write_text(
        """
        <html>
          <head><title>Fallback</title></head>
          <body>
            <div class="sidebar"><a href="/a">Nav A</a><a href="/b">Nav B</a><a href="/c">Nav C</a></div>
            <article class="prose">
              <h1>Deep Guide</h1>
              <p>This paragraph contains enough useful article text to win the content selection score.</p>
              <pre><code>print("hello")</code></pre>
            </article>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = convert_html(path)

    assert result.title == "Deep Guide"
    assert "useful article text" in result.body
    assert "```" in result.body
    assert "Nav A" not in result.body


def test_convert_pdf_empty_scanned_pdf_gets_warning(tmp_path: Path, monkeypatch):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF")

    class EmptyReader:
        pages = []

    monkeypatch.setattr("app.converters.documents.PdfReader", lambda _path: EmptyReader())

    result = convert_pdf(path)

    assert result.body == ""
    assert "OCR" in result.metadata["warning"]


def test_convert_pdf_uses_pymupdf_metadata_and_assets(tmp_path: Path):
    import fitz

    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Sample PDF Title", fontsize=20)
    page.insert_text((72, 110), "Abstract", fontsize=14)
    page.insert_text((72, 140), "This is extractable PDF text.", fontsize=11)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 16, 16), False)
    pixmap.clear_with(0x336699)
    page.insert_image(fitz.Rect(72, 180, 120, 228), pixmap=pixmap)
    doc.set_metadata({"title": "Sample PDF Title", "author": "Test Author"})
    doc.save(path)
    doc.close()

    result = convert_pdf(path, tmp_path / "assets")

    assert result.title == "Sample PDF Title"
    assert result.author == "Test Author"
    assert "### Abstract" in result.body
    assert "This is extractable PDF text." in result.body
    assert "![pdf-page-1-image-1](./assets/pdf-page-1-image-1.png)" in result.body
    assert "- 图片资源：./assets/pdf-page-1-image-1.png" in result.resources
    assert (tmp_path / "assets" / "pdf-page-1-image-1.png").exists()
    assert result.metadata["pdf_engine"] == "pymupdf"
    assert result.metadata["image_count"] == 1


def test_convert_pdf_infers_title_from_first_page_layout(tmp_path: Path):
    import fitz

    path = tmp_path / "paper-file-name.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Provided proper attribution is provided.", fontsize=11)
    page.insert_text((180, 150), "Attention Is All You Need", fontsize=18)
    page.insert_text((72, 220), "Abstract", fontsize=12)
    page.insert_text((72, 250), "This paper has extractable body text.", fontsize=10)
    doc.save(path)
    doc.close()

    result = convert_pdf(path, tmp_path / "assets")

    assert result.title == "Attention Is All You Need"
    assert "### Attention Is All You Need" not in result.body
    assert "### Abstract" in result.body
    assert "This paper has extractable body text." in result.body


def test_convert_pdf_extracts_detected_tables(tmp_path: Path):
    import fitz

    path = tmp_path / "table.pdf"
    doc = fitz.open()
    page = doc.new_page()
    x0, y0 = 72, 72
    for index in range(4):
        page.draw_line((x0, y0 + index * 30), (x0 + 240, y0 + index * 30))
    for index in range(3):
        page.draw_line((x0 + index * 120, y0), (x0 + index * 120, y0 + 90))
    for row_index, row in enumerate([["Name", "Value"], ["Alpha", "42"], ["Beta", "99"]]):
        for column_index, value in enumerate(row):
            page.insert_text((x0 + column_index * 120 + 8, y0 + row_index * 30 + 20), value, fontsize=10)
    doc.save(path)
    doc.close()

    result = convert_pdf(path, tmp_path / "assets")

    assert "| Name | Value |" in result.body
    assert "| Alpha | 42 |" in result.body
    assert result.metadata["table_count"] == 1
    assert result.metadata["row_count"] == 2


def test_sparse_pdf_tables_render_as_readable_form_regions():
    rows = [
        ["Step 1", "First name", "", "Last name", "", "OMB No. 1545-0074"],
        ["", "Address", "", "", "Long instructions that explain what the user should enter in this part of the form.", ""],
        ["Step 2", "", "Multiple Jobs", "", "", ""],
    ]

    assert should_render_pdf_table_as_form(rows)
    rendered = render_pdf_form_table(rows)

    assert rendered.startswith("**Form/table region**")
    assert "- Step 1" in rendered
    assert "| --- |" not in rendered


def test_collapsed_wide_pdf_tables_render_as_layout_blocks():
    rows = [
        ["", "train N d model d ff h d k d v P drop ls steps", "PPL BLEU params dev dev"],
        ["base", "6 512 2048 8 64 64 0.1 0.1 100K", "4.92 25.8 65"],
        ["A", "1 512 512 4 128 128 16 32 32 32 16 16", "5.29 24.9 5.00 25.5 4.91 25.8"],
    ]

    assert should_render_pdf_table_as_layout_block(rows)
    rendered = render_pdf_table_layout_block(rows)

    assert rendered.startswith("**Table region")
    assert "```text" in rendered
    assert "| --- |" not in rendered


def test_multiline_pdf_tables_expand_stacked_rows():
    rows = [
        ["Parser", "Training", "WSJ 23 F1"],
        [
            "Vinyals & Kaiser (2014)\nPetrov et al. (2006)",
            "WSJ only\nWSJ only",
            "88.3\n90.4",
        ],
        ["Transformer", "WSJ only", "91.3"],
    ]

    rendered = render_pdf_markdown_table(rows)

    assert "| Vinyals & Kaiser (2014) | WSJ only | 88.3 |" in rendered
    assert "| Petrov et al. (2006) | WSJ only | 90.4 |" in rendered
    assert "| Transformer | WSJ only | 91.3 |" in rendered


def test_multilevel_pdf_table_headers_are_not_split_into_data_rows():
    rows = [
        ["Higher Paying Job", "Lower Paying Job", "", "", ""],
        ["", "$0 -\n9,999", "$10,000 -\n19,999", "$20,000 -\n29,999", "$30,000 -\n39,999"],
        ["$0 - 9,999\n$10,000 - 19,999", "$0\n0", "$0\n480", "$480\n1,480", "$850\n1,850"],
    ]

    rendered = render_pdf_markdown_table(rows)

    assert "| Higher Paying Job | $0 - 9,999 | $10,000 - 19,999 |" in rendered
    assert "|  | 9,999 | 19,999 |" not in rendered
    assert "| $0 - 9,999 | $0 | $0 | $480 | $850 |" in rendered
    assert "| $10,000 - 19,999 | 0 | 480 | 1,480 | 1,850 |" in rendered


def test_coordinate_pdf_table_reconstructs_collapsed_scientific_columns():
    def word(x: float, y: float, text: str):
        return (x, y, x + max(len(text), 1) * 4.0, y + 7.0, text)

    words = [
        word(370, 100, "train"),
        word(405, 100, "PPL"),
        word(440, 100, "BLEU"),
        word(475, 100, "params"),
        word(146, 108, "N"),
        word(170, 108, "dmodel"),
        word(205, 108, "dff"),
        word(236, 108, "h"),
        word(258, 108, "dk"),
        word(285, 108, "dv"),
        word(310, 108, "Pdrop"),
        word(346, 108, "eps"),
        word(370, 116, "steps"),
        word(405, 116, "(dev)"),
        word(440, 116, "(dev)"),
        word(475, 116, "x106"),
        word(116, 126, "base"),
        word(148, 126, "6"),
        word(171, 126, "512"),
        word(202, 126, "2048"),
        word(237, 126, "8"),
        word(259, 126, "64"),
        word(285, 126, "64"),
        word(315, 126, "0.1"),
        word(345, 126, "0.1"),
        word(369, 126, "100K"),
        word(405, 126, "4.92"),
        word(440, 126, "25.8"),
        word(483, 126, "65"),
        word(237, 136, "1"),
        word(256, 136, "512"),
        word(283, 136, "512"),
        word(405, 136, "5.29"),
        word(440, 136, "24.9"),
        word(118, 142, "(A)"),
        word(237, 146, "4"),
        word(256, 146, "128"),
        word(283, 146, "128"),
        word(405, 146, "5.00"),
        word(440, 146, "25.5"),
    ]

    rows = reconstruct_pdf_table_from_word_items(words)
    rendered = render_coordinate_pdf_table(rows)

    assert len(rows[0]) >= 10
    assert "train steps" in rendered
    assert "| base | 6 | 512 | 2048 | 8 | 64 | 64 | 0.1 | 0.1 | 100K | 4.92 | 25.8 | 65 |" in rendered
    assert ["(A)", "", "", "", "1", "512", "512", "", "", "", "5.29", "24.9", ""] in rows


def test_coordinate_pdf_table_reuses_extracted_rowspan_labels():
    extracted_rows = [
        ["", "train N dmodel dff h dk dv Pdrop eps steps", "PPL BLEU params"],
        ["base", "6 512 2048 8 64 64 0.1 0.1 100K", "4.92 25.8 65"],
        ["(A)", "1 512 512\n4 128 128", "5.29 24.9\n5.00 25.5"],
        ["(B)", "16\n32", "5.16 25.1 58\n5.01 25.4 60"],
    ]
    coordinate_rows = [
        ["", "N", "dmodel", "dff", "h", "dk", "dv", "Pdrop", "eps", "steps", "PPL", "BLEU", "params"],
        ["base", "6", "512", "2048", "8", "64", "64", "0.1", "0.1", "100K", "4.92", "25.8", "65"],
        ["", "", "", "", "1", "512", "512", "", "", "", "5.29", "24.9", ""],
        ["(A)", "", "", "", "4", "128", "128", "", "", "", "5.00", "25.5", ""],
        ["", "", "", "", "", "16", "", "", "", "", "5.16", "25.1", "58"],
        ["(B)", "", "", "", "", "32", "", "", "", "", "5.01", "25.4", "60"],
    ]

    relabeled = apply_extracted_row_labels_to_coordinate_rows(extracted_rows, coordinate_rows)

    assert relabeled[2][0] == "(A)"
    assert relabeled[3][0] == ""
    assert relabeled[4][0] == "(B)"
    assert relabeled[5][0] == ""


def test_coordinate_pdf_table_collapses_empty_spanner_columns():
    rows = [
        ["Benchmark", "GPT-4", "GPT-3.5", "Contamination", "GPT-4 (non-contaminated)", "", "", "Degradation"],
        ["MMLU", "86.4%", "70.0%", "~0.6%", "", "", "-", "-"],
        ["HumanEval", "67.0%", "48.1%", "25%", "", "65.58%", "", "-2.12%"],
    ]

    collapsed = collapse_pdf_spanner_columns(rows)
    rendered = render_coordinate_pdf_table(rows)

    assert collapsed[0] == ["Benchmark", "GPT-4", "GPT-3.5", "Contamination", "GPT-4 (non-contaminated)", "Degradation"]
    assert collapsed[1] == ["MMLU", "86.4%", "70.0%", "~0.6%", "-", "-"]
    assert collapsed[2] == ["HumanEval", "67.0%", "48.1%", "25%", "65.58%", "-2.12%"]
    assert "| Benchmark | GPT-4 | GPT-3.5 | Contamination | GPT-4 (non-contaminated) | Degradation |" in rendered


def test_coordinate_pdf_table_does_not_take_over_two_column_text_tables():
    extracted_rows = [
        ["English", "Swahili"],
        [
            "A highly knowledgeable artificial intelligence model\nanswers multiple-choice questions\nA) Lower variance\nB) Higher variance",
            "Muundo wa akili bandia wenye ujuzi\nhujibu maswali ya chaguo-nyingi\nA) Tofauti ya chini\nB) Tofauti ya juu\nAnswer:",
        ],
    ]
    coordinate_rows = [
        ["English", "", "", "", "", "Swahili", "", "", "", ""],
        ["A", "highly", "knowledgeable", "model", "", "Muundo", "wa", "akili", "bandia", "wenye"],
        ["answers", "multiple-choice", "questions", "", "", "hujibu", "maswali", "ya", "chaguo-nyingi", ""],
        ["A)", "Lower", "variance", "", "", "A)", "Tofauti", "ya", "chini", ""],
        ["B)", "Higher", "variance", "", "", "B)", "Tofauti", "ya", "juu", ""],
    ]

    assert should_prefer_extracted_pdf_table(extracted_rows, coordinate_rows)
    assert not should_use_coordinate_pdf_table(extracted_rows, coordinate_rows)
    rendered = render_pdf_markdown_table(extracted_rows)

    assert "| English | Swahili |" in rendered
    assert "| A highly knowledgeable artificial intelligence model | Muundo wa akili bandia wenye ujuzi |" in rendered
    assert "| B) Higher variance | B) Tofauti ya juu |" in rendered
    assert "|  | Answer: |" in rendered
    assert "| highly |" not in rendered


def test_paragraph_like_pdf_table_false_positive_is_rejected():
    rows = [
        ["Section 4.4 requires CSPs to use measures to maintain the objectives of", "predictability (enabling"],
        ["reliable assumptions by individuals, owners, and operators about PII and its processing by an", ""],
        ["information system) and manageability (providing the capability for granular administration of", ""],
        ["PII, including alteration, deletion, and selective disclosure)", ""],
    ]

    assert is_false_positive_pdf_paragraph_table(rows)


def test_grid_pdf_table_reconstructs_cells_from_line_boundaries():
    def word(x: float, y: float, text: str):
        return (x, y, x + max(len(text), 1) * 4.0, y + 7.0, text)

    words = [
        word(80, 82, "Name"),
        word(202, 82, "Value"),
        word(80, 112, "Alpha"),
        word(202, 112, "42"),
        word(80, 142, "Beta"),
        word(202, 142, "99"),
    ]
    segments = [
        PdfLineSegment("h", 72, 72, 312),
        PdfLineSegment("h", 102, 72, 312),
        PdfLineSegment("h", 132, 72, 312),
        PdfLineSegment("h", 162, 72, 312),
        PdfLineSegment("v", 72, 72, 162),
        PdfLineSegment("v", 192, 72, 162),
        PdfLineSegment("v", 312, 72, 162),
    ]

    rows = reconstruct_pdf_grid_table(words, segments, bbox=(72, 72, 312, 162))
    rendered = render_coordinate_pdf_table(rows)

    assert rows == [["Name", "Value"], ["Alpha", "42"], ["Beta", "99"]]
    assert "| Name | Value |" in rendered
    assert "| Alpha | 42 |" in rendered


def test_sparse_visual_pdf_tables_render_as_layout_blocks():
    rows = [["" for _ in range(16)] for _ in range(5)]
    rows[1][0] = "The"
    rows[1][2] = "Law"
    rows[1][4] = "will"
    rows[2][8] = "attention"
    rows[3][14] = "heads"

    assert should_render_pdf_table_as_layout_block(rows)
    rendered = render_pdf_table_layout_block(rows)

    assert rendered.startswith("**Table region")
    assert "| --- |" not in rendered


def test_convert_pdf_falls_back_to_pypdf(monkeypatch, tmp_path: Path):
    path = tmp_path / "fallback.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF")

    class Page:
        def extract_text(self):
            return "Abstract\nFallback text."

    class Reader:
        pages = [Page()]
        metadata = {"/Title": "Fallback PDF"}

    monkeypatch.setattr("app.converters.documents.PdfReader", lambda _path: Reader())
    monkeypatch.setattr("app.converters.documents.extract_pdf_with_pymupdf", lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("boom")))

    result = convert_pdf(path, tmp_path / "assets")

    assert result.title == "Fallback PDF"
    assert result.metadata["pdf_engine"] == "pypdf"
    assert result.metadata["pdf_engine_fallback_reason"] == "boom"
    assert "Fallback text." in result.body


def test_pdf_text_normalization_repairs_lines_and_detects_headings():
    text = normalize_pdf_text(
        "INTRODUCTION\n"
        "This is a para-\n"
        "graph split across\n"
        "several extraction lines.\n\n"
        "1.1 Key idea\n"
        "- bullet item\n"
    )

    assert "### INTRODUCTION" in text
    assert "paragraph split across several extraction lines." in text
    assert "### 1.1 Key idea" in text
    assert "- bullet item" in text


def test_pdf_block_rendering_detects_layout_heading_caption_and_list():
    blocks = [
        PdfTextBlock(page=1, text="Large Section", bbox=(72, 80, 300, 100), font_size=16, is_bold=True),
        PdfTextBlock(page=1, text="Figure 1. A useful diagram.", bbox=(72, 120, 300, 140), font_size=9),
        PdfTextBlock(page=1, text="1) ordered item", bbox=(72, 160, 300, 180), font_size=10),
        PdfTextBlock(page=1, text="Normal paragraph continues", bbox=(72, 200, 300, 220), font_size=10),
        PdfTextBlock(page=1, text="until it ends.", bbox=(72, 225, 300, 245), font_size=10),
    ]

    rendered = render_pdf_blocks(blocks, body_size=10)

    assert "### Large Section" in rendered
    assert "> Figure 1. A useful diagram." in rendered
    assert "1) ordered item" in rendered
    assert "Normal paragraph continues until it ends." in rendered
    assert is_layout_pdf_heading("Large Section", blocks[0], body_size=10)
    assert not is_layout_pdf_heading("Paul A. Grassi", PdfTextBlock(page=1, text="Paul A. Grassi", bbox=(72, 80, 200, 100), font_size=16), body_size=10)


def test_pdf_block_rendering_preserves_formula_blocks():
    blocks = [
        PdfTextBlock(
            page=1,
            text="z(x) = ReLU(Wencx + benc),",
            bbox=(100, 100, 300, 120),
            font_size=10,
            is_formula=True,
        ),
        PdfTextBlock(page=1, text="(1)", bbox=(280, 100, 300, 120), font_size=10, is_formula=True),
    ]

    rendered = render_pdf_blocks(blocks, body_size=10)

    assert "$$\nz(x) = ReLU(Wencx + benc), (1)\n$$" in rendered
    assert "### z(x)" not in rendered


def test_pdf_block_rendering_keeps_explanations_outside_formula_blocks():
    blocks = [
        PdfTextBlock(page=3, text="ws→t = f (s) ⊤\ndec\nJ▼\n(s)→(t) f (t)\nenc, (6)", bbox=(80, 80, 300, 130), font_size=10, is_formula=True),
        PdfTextBlock(page=3, text="where f (s)\ndec ∈ Rdmodel is the decoder vector.", bbox=(80, 140, 300, 170), font_size=10),
        PdfTextBlock(page=3, text="3", bbox=(300, 760, 310, 770), font_size=10),
    ]

    rendered = render_pdf_blocks(blocks, body_size=10)
    formula = rendered.split("$$")[1]

    assert "ws→t = f (s)" in formula
    assert "where f (s)" not in formula
    assert "where f (s) dec ∈ Rdmodel is the decoder vector." in rendered
    assert "\n3\n" not in f"\n{rendered}\n"


def test_pdf_block_rendering_does_not_promote_numbered_body_lists_to_headings():
    blocks = [
        PdfTextBlock(
            page=2,
            text="1. Hierarchical Integration of Visual and Semantic Con-\ncepts: Features emerge only in higher layers.",
            bbox=(80, 200, 300, 240),
            font_size=10,
        )
    ]

    rendered = render_pdf_blocks(blocks, body_size=10)

    assert "### 1. Hierarchical" not in rendered
    assert "1. Hierarchical Integration" in rendered


def test_convert_docx_with_heading_table_and_image_dir(tmp_path: Path):
    path = tmp_path / "report.docx"
    document = Document()
    document.add_heading("Report Title", level=1)
    document.add_paragraph("A useful paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    document.save(path)

    result = convert_docx(path, tmp_path / "assets")

    assert result.title == "Report Title"
    rendered = render_document_markdown(title=result.title, source_type=result.source_type, body=result.body, summary="")
    assert rendered.count("# Report Title") == 1
    assert "A useful paragraph." in result.body
    assert "| Name | Value |" in result.body
    assert "| Alpha | 42 |" in result.body
    assert result.body.index("A useful paragraph.") < result.body.index("| Name | Value |")
    assert result.metadata["heading_count"] == 1
    assert result.metadata["table_count"] == 1


def test_convert_docx_preserves_inline_formatting(tmp_path: Path):
    path = tmp_path / "formatted.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Plain ")
    bold = paragraph.add_run("bold")
    bold.bold = True
    paragraph.add_run(" and ")
    italic = paragraph.add_run("italic")
    italic.italic = True
    document.save(path)

    result = convert_docx(path, tmp_path / "assets")

    assert "Plain **bold** and *italic*" in result.body
